#!/usr/bin/env python3
"""
知识库网页服务 — 基于 Hindsight + DeepSeek
文档上传 → 自动切片入库 → 自然语言搜索 → AI 合成答案
v2: 增加标题/分类系统，本地 SQLite 存储文档元数据
"""
import os, sys, json, re, uuid, sqlite3, subprocess, tempfile, shutil, traceback, hashlib, asyncio
from pathlib import Path
from datetime import datetime
from io import BytesIO
import zipfile as _zipfile
import time as _time

import httpx
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
import pypdf
import docx
from rank_bm25 import BM25Okapi
import jieba

# ─── 环境 ────────────────────────────────────────────────────
ENV_FILE = os.path.expanduser("~/.hermes/.env")
if os.path.exists(ENV_FILE):
    for line in open(ENV_FILE):
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, v = line.split("=", 1)
            os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))

DEEPSEEK_KEY = os.environ.get("XIAOMI_API_KEY", "")
DEEPSEEK_BASE = os.environ.get("XIAOMI_BASE_URL", "https://token-plan-cn.xiaomimimo.com/v1")
DEEPSEEK_MODEL = "mimo-v2.5"
HINDSIGHT_URL = "http://localhost:8888"
KB_BANK = "kb"  # 默认 bank（兼容旧数据，"全部"查询时用）

# ── 多 Bank 配置 ──
BANKS = {
    "all":           {"name": "全部",           "hindsight": "kb", "prompt": "通用政务信息化知识库"},
    "project_docs":  {"name": "项目资料",       "hindsight": "kb", "prompt": "你是政务信息化项目管理专家。熟悉项目管理办法、验收管理细则、财政投资规定、软件行业基准数据。回答时注重管理流程、审批要求和实操经验。"},
    "standards":     {"name": "规范",           "hindsight": "kb", "prompt": "你是政务信息化标准规范专家。精通GB/GA/T/EGAG/GDZW等国家及团体标准，覆盖等保测评、密码应用、监理服务、立项咨询、验收测评、会议系统、安防工程、数据中心等领域。回答时注重条款引用和合规要求。"},
    "industry_docs": {"name": "信息化行业文档", "hindsight": "kb", "prompt": "你是政务信息化行业专家。熟悉电子政务工程造价、软件造价评估、信创替代、验收测评实务、行业政策解读。回答时注重实操经验和行业惯例。"},
    "templates":     {"name": "方案模板",       "hindsight": "kb", "prompt": "你是政务信息化项目方案编写专家。精通建设开发类和运维服务类项目方案的编写规范、章节结构、技术路线选型。回答时注重模板结构和编写要点。"},
    "tech_guides":   {"name": "技术指导书",     "hindsight": "kb", "prompt": "你是全栈技术专家。精通前端/后端/Agent/DevOps/安全/渗透测试/AI/LLM。回答注重实战经验、架构设计和攻防思路。"},
    "general":       {"name": "综合文件",       "hindsight": "kb", "prompt": "你是知识管理助手。擅长整理归纳各类知识，回答清晰有条理。"},
}

# MinerU API 配置
MINERU_TOKEN = os.environ.get("MINERU_API_TOKEN", "")
MINERU_BASE = "https://mineru.net/api/v4"
MINERU_PAGES_MAX = 200  # 单次API调用的最大页数

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "meta.db")

app = FastAPI(title="知识库", docs_url=None)

# ─── 默认分类列表 ────────────────────────────────────────────────
DEFAULT_CATEGORIES = [
    "💡想法", "💼工作", "📚学习", "🏠生活", "🚀项目",
    "💭灵感", "📝会议", "🔧技术", "📊数据", "📰资讯",
    "🔒安全", "🤖AI", "其他",
]

# ─── SQLite 元数据 ──────────────────────────────────────────────
def get_db():
    db = sqlite3.connect(DB_PATH)
    db.row_factory = sqlite3.Row
    db.execute("PRAGMA journal_mode=WAL")
    return db

def init_db():
    db = get_db()
    db.execute("""
        CREATE TABLE IF NOT EXISTS doc_meta (
            doc_id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            category TEXT DEFAULT '',
            filename TEXT DEFAULT '',
            content_hash TEXT DEFAULT '',
            created_at TEXT NOT NULL
        )
    """)
    # 兼容旧表：补加 content_hash 列
    try:
        db.execute("ALTER TABLE doc_meta ADD COLUMN content_hash TEXT DEFAULT ''")
        print("DB: 已添加 content_hash 列", file=sys.stderr)
    except sqlite3.OperationalError:
        pass  # 列已存在
    # doc_type 列（adaptive chunking 策略类型）
    try:
        db.execute("ALTER TABLE doc_meta ADD COLUMN doc_type TEXT DEFAULT 'generic'")
        print("DB: 已添加 doc_type 列", file=sys.stderr)
    except sqlite3.OperationalError:
        pass  # 列已存在
    # Parent-Child 分块表
    db.execute("""CREATE TABLE IF NOT EXISTS parent_chunks (
        doc_id TEXT,
        parent_idx INTEGER,
        parent_text TEXT,
        PRIMARY KEY (doc_id, parent_idx)
    )""")
    db.commit()
    db.close()

init_db()

def save_meta(doc_id: str, title: str, category: str, filename: str, content_hash: str = "", doc_type: str = "generic"):
    db = get_db()
    db.execute(
        "INSERT OR REPLACE INTO doc_meta (doc_id, title, category, filename, content_hash, created_at, doc_type) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (doc_id, title, category, filename, content_hash, datetime.utcnow().isoformat(), doc_type),
    )
    db.commit()
    db.close()

def get_meta(doc_id: str) -> dict:
    db = get_db()
    row = db.execute("SELECT * FROM doc_meta WHERE doc_id = ?", (doc_id,)).fetchone()
    db.close()
    if row:
        return {"title": row["title"], "category": row["category"], "filename": row["filename"], "created_at": row["created_at"]}
    return {"title": "未知文档", "category": "", "filename": "未知", "created_at": ""}

def get_all_meta() -> dict:
    """返回 {doc_id: {...}}"""
    db = get_db()
    rows = db.execute("SELECT * FROM doc_meta ORDER BY created_at DESC").fetchall()
    db.close()
    result = {}
    for r in rows:
        result[r["doc_id"]] = {"title": r["title"], "category": r["category"], "filename": r["filename"], "created_at": r["created_at"]}
    return result

def update_meta(doc_id: str, title: str = None, category: str = None):
    db = get_db()
    if title is not None and category is not None:
        db.execute("UPDATE doc_meta SET title=?, category=? WHERE doc_id=?", (title, category, doc_id))
    elif title is not None:
        db.execute("UPDATE doc_meta SET title=? WHERE doc_id=?", (title, doc_id))
    elif category is not None:
        db.execute("UPDATE doc_meta SET category=? WHERE doc_id=?", (category, doc_id))
    db.commit()
    db.close()

def delete_meta(doc_id: str):
    db = get_db()
    db.execute("DELETE FROM doc_meta WHERE doc_id = ?", (doc_id,))
    db.commit()
    db.close()

def find_by_hash(content_hash: str) -> dict | None:
    """按内容哈希查找已存在的文档，返回元数据或 None"""
    if not content_hash:
        return None
    db = get_db()
    row = db.execute(
        "SELECT doc_id, title, category, filename, created_at FROM doc_meta WHERE content_hash = ? AND content_hash != '' AND bank != 'skip'",
        (content_hash,)
    ).fetchone()
    db.close()
    if row:
        return {
            "doc_id": row["doc_id"],
            "title": row["title"],
            "category": row["category"],
            "filename": row["filename"],
            "created_at": row["created_at"],
        }
    return None

# ─── 工具函数 ──────────────────────────────────────────────────

MAX_FILE_SIZE = 200 * 1024 * 1024  # 200MB（与 MinerU 对齐）

def ocr_pdf(pdf_bytes: bytes) -> str:
    """用 pdftoppm + tesseract 对扫描件 PDF 做 OCR，返回纯文本"""
    tmpdir = tempfile.mkdtemp(prefix="kb_ocr_")
    try:
        pdf_path = os.path.join(tmpdir, "input.pdf")
        with open(pdf_path, "wb") as f:
            f.write(pdf_bytes)

        # 先探测页数
        page_count = len(pypdf.PdfReader(BytesIO(pdf_bytes)).pages)
        print(f"OCR: {page_count} 页扫描件，开始转换...", file=sys.stderr)

        # PDF → 灰度 PNG（200 DPI 平衡速度与精度）
        try:
            result = subprocess.run(
                ["pdftoppm", "-png", "-gray", "-r", "200", pdf_path, os.path.join(tmpdir, "page")],
                capture_output=True, text=True, timeout=600,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                f"PDF 转换图片超时（{page_count} 页扫描件处理超过 5 分钟）。"
                "建议先用 PC 端工具将 PDF 压缩/降低分辨率后再上传。"
            )
        except Exception as e:
            traceback.print_exc(file=sys.stderr)
            raise RuntimeError(f"PDF 转换异常: {e}")

        if result.returncode != 0:
            raise RuntimeError(f"pdftoppm 失败: {result.stderr[:200]}")

        # 逐页 OCR
        pages = sorted(Path(tmpdir).glob("page-*.png"))
        if not pages:
            raise RuntimeError("PDF 转换后无图片输出")

        texts = []
        for idx, png in enumerate(pages):
            if idx % 5 == 0 and idx > 0:
                print(f"OCR: {idx}/{len(pages)} 页已完成...", file=sys.stderr)
            out_base = os.path.join(tmpdir, f"ocr_{png.stem}")
            try:
                ocr_result = subprocess.run(
                    ["tesseract", str(png), out_base, "-l", "chi_sim+eng", "--psm", "3"],
                    capture_output=True, text=True, timeout=60,
                )
            except subprocess.TimeoutExpired:
                print(f"OCR 页面 {png.name} 超时，跳过", file=sys.stderr)
                continue
            except Exception as e:
                print(f"OCR 页面 {png.name} 异常: {e}", file=sys.stderr)
                continue
            if ocr_result.returncode != 0:
                print(f"OCR 页面 {png.name} 失败: {ocr_result.stderr[:100]}", file=sys.stderr)
                continue
            out_txt = out_base + ".txt"
            if os.path.exists(out_txt):
                with open(out_txt, "r", encoding="utf-8") as f:
                    texts.append(f.read().strip())
            os.unlink(out_txt)

        return "\n\n".join(texts)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

async def mineru_parse_pdf(filename: str, content: bytes) -> str:
    """通过 MinerU API 解析 PDF，返回 Markdown（含HTML表格）。
    
    流程：获取上传URL → PUT文件 → 轮询结果 → 下载ZIP → 提取 full.md
    
    超过 MINERU_PAGES_MAX 页自动分批，合并结果。
    """
    if not MINERU_TOKEN:
        raise RuntimeError("MINERU_API_TOKEN 未配置")
    
    reader = pypdf.PdfReader(BytesIO(content))
    total_pages = len(reader.pages)
    print(f"MinerU: {total_pages} 页，开始解析...", file=sys.stderr)
    
    # 计算分批范围
    ranges = []
    for start in range(0, total_pages, MINERU_PAGES_MAX):
        if start >= total_pages:
            break
        end = min(start + MINERU_PAGES_MAX, total_pages)
        ranges.append((start + 1, end))  # page_ranges 是 1-indexed
    
    print(f"MinerU: 分 {len(ranges)} 批: {ranges}", file=sys.stderr)
    
    all_md_parts = []
    
    for batch_idx, (pg_start, pg_end) in enumerate(ranges):
        batch_label = f"batch{batch_idx+1}"
        page_range = f"{pg_start}-{pg_end}"
        
        # Step 1: 获取上传 URL
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(
                f"{MINERU_BASE}/file-urls/batch",
                headers={
                    "Authorization": f"Bearer {MINERU_TOKEN}",
                    "Content-Type": "application/json",
                },
                json={
                    "files": [{
                        "name": f"{batch_label}.pdf",
                        "is_ocr": True,
                        "page_ranges": page_range,
                    }],
                    "model_version": "pipeline",
                    "language": "ch",
                    "enable_table": True,
                },
            )
            data = resp.json()
            if data.get("code") != 0:
                raise RuntimeError(f"MinerU 获取上传URL失败: {data.get('msg')}")
            
            batch_id = data["data"]["batch_id"]
            file_url = data["data"]["file_urls"][0]
        
        # Step 2: 上传文件
        # 如果文件超过200页需要截取，否则传整个文件
        if len(ranges) == 1:
            upload_bytes = content
        else:
            # 截取指定页码范围
            writer = pypdf.PdfWriter()
            for pg in range(pg_start - 1, pg_end):
                writer.add_page(reader.pages[pg])
            upload_buf = BytesIO()
            writer.write(upload_buf)
            upload_bytes = upload_buf.getvalue()
        
        async with httpx.AsyncClient(timeout=120) as client:
            resp = await client.put(file_url, content=upload_bytes)
            if resp.status_code != 200:
                raise RuntimeError(f"MinerU 文件上传失败: HTTP {resp.status_code}")
        
        print(f"MinerU {batch_label}: 已上传 {pg_start}-{pg_end} 页", file=sys.stderr)
        
        # Step 3: 轮询结果
        poll_url = f"{MINERU_BASE}/extract-results/batch/{batch_id}"
        max_wait = 600  # 最多等10分钟
        
        for _ in range(max_wait // 3):
            await _asleep(3)
            async with httpx.AsyncClient(timeout=30) as client:
                resp = await client.get(
                    poll_url,
                    headers={"Authorization": f"Bearer {MINERU_TOKEN}"},
                )
                data = resp.json()
            
            if data.get("code") != 0:
                continue
            
            er = (data.get("data", {}).get("extract_result") or [{}])[0]
            state = er.get("state")
            
            if state == "done":
                zip_url = er.get("full_zip_url")
                break
            elif state == "failed":
                raise RuntimeError(f"MinerU {batch_label} 失败: {er.get('err_msg')}")
        else:
            raise RuntimeError(f"MinerU {batch_label}: 轮询超时 ({max_wait}s)")
        
        # Step 4: 下载并解压 Markdown
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.get(zip_url)
            zip_data = resp.content
        
        with _zipfile.ZipFile(BytesIO(zip_data)) as zf:
            if "full.md" not in zf.namelist():
                raise RuntimeError(f"MinerU {batch_label}: ZIP 中缺少 full.md")
            md_text = zf.read("full.md").decode("utf-8")
        
        all_md_parts.append(md_text)
        print(f"MinerU {batch_label}: {len(md_text)} 字符", file=sys.stderr)
    
    result = "\n\n".join(all_md_parts)
    print(f"MinerU 完成: {len(result)} 字符", file=sys.stderr)
    return result


async def docx_to_pdf_via_libreoffice(filename: str, content: bytes) -> bytes:
    """用 LibreOffice headless 将 DOCX 转为 PDF，返回 PDF 字节"""
    import tempfile
    tmpdir = tempfile.mkdtemp(prefix="kb_docx2pdf_")
    try:
        docx_path = os.path.join(tmpdir, filename)
        with open(docx_path, "wb") as f:
            f.write(content)
        
        proc = await asyncio.create_subprocess_exec(
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", tmpdir, docx_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=120)
        
        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice 退出码 {proc.returncode}: {stderr.decode()[:200]}")
        
        pdf_name = os.path.splitext(filename)[0] + ".pdf"
        pdf_path = os.path.join(tmpdir, pdf_name)
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"PDF 未生成: {pdf_path}")
        
        with open(pdf_path, "rb") as f:
            return f.read()
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


async def _asleep(seconds: float):
    """异步 sleep，兼容 Python 3.11-"""
    import asyncio
    await asyncio.sleep(seconds)

async def parse_document(filename: str, content: bytes) -> str:
    """解析 PDF/Word/Markdown/TXT → 纯文本
    
    PDF 优先使用 MinerU API（高精度表格识别），失败时回退到 tesseract OCR。
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        try:
            reader = pypdf.PdfReader(BytesIO(content))
        except Exception as e:
            raise ValueError(f"PDF 解析失败（文件可能已损坏或加密）: {e}")
        if reader.is_encrypted:
            raise ValueError("PDF 文件已加密，无法提取文字内容")
        
        text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
        if not text.strip():
            page_count = len(reader.pages)
            print(f"PDF 文字层为空 ({page_count} 页)，启动 OCR...", file=sys.stderr)
            
            # 优先尝试 MinerU API
            if MINERU_TOKEN:
                try:
                    text = await mineru_parse_pdf(filename, content)
                    if text.strip():
                        print(f"MinerU 完成，提取 {len(text)} 字符", file=sys.stderr)
                        return text
                except Exception as e:
                    print(f"MinerU 失败，回退 tesseract: {e}", file=sys.stderr)
                    traceback.print_exc(file=sys.stderr)
            
            # 回退到 tesseract
            print(f"使用 tesseract OCR ({page_count} 页)...", file=sys.stderr)
            try:
                text = ocr_pdf(content)
                if not text.strip():
                    raise ValueError(
                        "PDF OCR 识别结果为空。"
                        "可能原因：①图片质量过低 ②PDF 为纯图片且文字不清晰。"
                    )
                print(f"tesseract 完成，提取 {len(text)} 字符", file=sys.stderr)
            except RuntimeError as e:
                raise ValueError(f"PDF OCR 失败: {e}")
            except Exception as e:
                traceback.print_exc(file=sys.stderr)
                raise ValueError(f"PDF OCR 异常: {e}")
        return text
    elif ext in (".docx", ".doc"):
        try:
            d = docx.Document(BytesIO(content))
        except Exception as e:
            raise ValueError(f"Word 文档解析失败: {e}")
        text = "\n\n".join(p.text for p in d.paragraphs)
        # 同时提取表格文本
        for table in d.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    text += "\n" + row_text
        if text.strip():
            return text
        # python-docx 解析为空（纯表格/文本框等）→ 转 PDF 再解析
        print(f"DOCX 段落+表格解析为空，尝试 LibreOffice 转 PDF: {filename}", file=sys.stderr)
        try:
            pdf_bytes = await docx_to_pdf_via_libreoffice(filename, content)
            if pdf_bytes:
                return await parse_document(
                    filename.replace(".docx", ".pdf").replace(".doc", ".pdf"),
                    pdf_bytes
                )
        except Exception as e:
            print(f"LibreOffice 转换失败: {e}", file=sys.stderr)
        raise ValueError("Word 文档内容为空，且自动转 PDF 也未能提取内容。请检查文件是否包含可读文字。")
    elif ext in (".md", ".markdown", ".txt", ".text", ""):
        return content.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

def filename_to_title(filename: str) -> str:
    """从文件名生成默认标题：去掉扩展名"""
    return Path(filename).stem or filename

async def deepseek_chat(messages: list, stream: bool = False) -> str:
    """调用 DeepSeek Chat API"""
    async with httpx.AsyncClient(timeout=120) as client:
        resp = await client.post(
            f"{DEEPSEEK_BASE}/chat/completions",
            headers={"Authorization": f"Bearer {DEEPSEEK_KEY}"},
            json={
                "model": DEEPSEEK_MODEL,
                "messages": messages,
                "temperature": 0.3,
                "max_tokens": 3000,
                "stream": stream,
            },
            timeout=120,
        )
        if stream:
            return resp
        data = resp.json()
        return data["choices"][0]["message"]["content"]

async def hindsight_request(endpoint: str, method: str = "GET", json_data: dict = None, timeout: int = 30) -> dict:
    """调用 Hindsight API"""
    async with httpx.AsyncClient(timeout=timeout) as client:
        try:
            if method == "POST":
                resp = await client.post(f"{HINDSIGHT_URL}{endpoint}", json=json_data)
            elif method == "DELETE":
                resp = await client.delete(f"{HINDSIGHT_URL}{endpoint}")
            else:
                resp = await client.get(f"{HINDSIGHT_URL}{endpoint}")
        except httpx.TimeoutException:
            raise Exception(f"Hindsight {method} {endpoint}: 请求超时（{timeout}s）")
        except httpx.ConnectError:
            raise Exception(f"Hindsight {method} {endpoint}: 无法连接（服务未启动？）")
        except Exception as e:
            raise Exception(f"Hindsight {method} {endpoint}: 网络异常: {e}")

        if resp.status_code >= 400:
            detail = ""
            try:
                detail = resp.json().get("detail", "")
            except Exception:
                pass
            detail = detail or resp.text[:200] or f"HTTP {resp.status_code}"
            raise Exception(f"Hindsight {method} {endpoint} returned {resp.status_code}: {detail}")
        try:
            return resp.json()
        except Exception:
            raise Exception(f"Hindsight {method} {endpoint}: 响应不是有效 JSON: {resp.text[:200]}")


# ── BM25 索引管理（带 TTL 缓存）──
_bm25_cache = {"index": None, "docs": [], "bank": None, "ts": 0}
_BM25_TTL = 300  # 5 分钟缓存

def _tokenize(text: str) -> list:
    """jieba 分词，去掉单字符和空白"""
    return [w for w in jieba.cut(text) if len(w.strip()) > 1]

async def build_bm25_index(bank: str = "all") -> tuple:
    """通过多组 recall 查询获取全部 chunks 构建 BM25 索引（带缓存）"""
    import time
    now = time.time()
    if _bm25_cache["index"] and _bm25_cache["bank"] == bank and (now - _bm25_cache["ts"]) < _BM25_TTL:
        return _bm25_cache["index"], _bm25_cache["docs"]

    docs = []
    # 使用多组通用查询并行召回，覆盖更多 chunks
    recall_queries = [
        "标准 规范", "安全 系统", "工程 技术", "设计 施工",
        "检测 验收", "网络 安全", "信息 系统", "监控 设备",
    ]
    seen_texts = set()
    try:
        tasks = [recall(q, limit=300, bank="kb", max_tokens=65536) for q in recall_queries]
        results_list = await asyncio.gather(*tasks, return_exceptions=True)
        for results in results_list:
            if isinstance(results, Exception) or not results:
                continue
            for r in results:
                text = r.get("text", "") or ""
                if not text.strip():
                    continue
                # 去重：取前 80 字符作为 key
                dedup_key = text[:80]
                if dedup_key in seen_texts:
                    continue
                seen_texts.add(dedup_key)
                tags = r.get("tags", [])
                doc_id = None
                for t in tags:
                    if t.startswith("doc_id:"):
                        doc_id = t[7:]
                        break
                docs.append({"text": text, "doc_id": doc_id or "_unknown_", "tags": tags})
    except Exception as e:
        print(f"[WARN] BM25 recall failed: {e}", flush=True)

    if not docs:
        return None, []

    tokenized = [_tokenize(d["text"]) for d in docs]
    bm25 = BM25Okapi(tokenized)

    _bm25_cache.update({"index": bm25, "docs": docs, "bank": bank, "ts": now})
    print(f"BM25 index built: {len(docs)} chunks (from {len(recall_queries)} queries)", flush=True)
    return bm25, docs

def bm25_search(query: str, bm25, docs: list, top_k: int = 10) -> list:
    """BM25 关键词搜索"""
    if not bm25 or not docs:
        return []
    query_tokens = _tokenize(query)
    scores = bm25.get_scores(query_tokens)
    # 取 top_k
    top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:top_k]
    results = []
    for idx in top_indices:
        if scores[idx] > 0:
            results.append({"text": docs[idx]["text"], "doc_id": docs[idx]["doc_id"], 
                          "tags": docs[idx]["tags"], "bm25_score": float(scores[idx])})
    return results

def rrf_merge(dense_results: list, bm25_results: list, k: int = 60) -> list:
    """Reciprocal Rank Fusion 融合两路召回结果"""
    doc_scores = {}
    doc_data = {}
    
    # Dense 结果按排名打分
    for rank, r in enumerate(dense_results):
        text = r.get("text", "")
        doc_id = None
        for t in r.get("tags", []):
            if t.startswith("doc_id:"):
                doc_id = t[7:]
                break
        key = doc_id or text[:50]
        doc_scores[key] = doc_scores.get(key, 0) + 1.0 / (k + rank + 1)
        if key not in doc_data:
            doc_data[key] = r
    
    # BM25 结果按排名打分
    for rank, r in enumerate(bm25_results):
        text = r.get("text", "")
        doc_id = r.get("doc_id")
        key = doc_id or text[:50]
        doc_scores[key] = doc_scores.get(key, 0) + 1.0 / (k + rank + 1)
        if key not in doc_data:
            doc_data[key] = r
    
    # 按 RRF 分数排序
    sorted_keys = sorted(doc_scores.keys(), key=lambda x: doc_scores[x], reverse=True)
    return [doc_data[k] for k in sorted_keys]

async def recall(query: str, limit: int = 5, bank: str = "kb", max_tokens: int = 4096) -> list:
    """语义召回 — 支持指定 bank"""
    result = await hindsight_request(
        f"/v1/default/banks/{bank}/memories/recall",
        "POST",
        {"query": query, "max_tokens": max_tokens},
    )
    return result.get("results", [])

def get_bank_config(bank_key: str) -> dict:
    """获取 bank 配置，不存在则返回默认"""
    return BANKS.get(bank_key, BANKS["all"])

def assess_quality(text: str) -> dict:
    """评估文本质量，返回 {score, total_chars, meaningful_chars, issues}"""
    if not text or len(text.strip()) < 50:
        return {"score": 0, "total_chars": len(text), "meaningful_chars": 0,
                "issues": ["文本过短（<50字符）"]}

    total = len(text)
    meaningful = 0
    garbage_chars = 0  # replacement char U+FFFD
    repeated_runs = 0

    for i, ch in enumerate(text):
        code = ord(ch)
        # CJK Unified Ideographs + Extension A
        if 0x4E00 <= code <= 0x9FFF or 0x3400 <= code <= 0x4DBF:
            meaningful += 1
        # ASCII letters/digits
        elif (0x30 <= code <= 0x39) or (0x41 <= code <= 0x5A) or (0x61 <= code <= 0x7A):
            meaningful += 1
        # Common punctuation
        elif 0x20 <= code <= 0x2F or 0x3A <= code <= 0x40:
            meaningful += 1
        elif code in (0x3001, 0x3002, 0xFF0C, 0xFF0E, 0xFF1A, 0xFF1B, 0xFF08, 0xFF09, 0x0A):
            meaningful += 1
        # Replacement character
        elif code == 0xFFFD:
            garbage_chars += 1
        # Repeated char detection (3+ same in a row, skip spaces/newlines)
        if i >= 2 and text[i] == text[i-1] == text[i-2] and ord(text[i]) > 32 and ord(text[i]) != 0x0A:
            repeated_runs += 1

    # Scores
    garbage_ratio = garbage_chars / max(total, 1)
    meaningful_ratio = meaningful / max(total, 1)
    repeated_ratio = repeated_runs / max(total, 1)

    # Weighted score
    score = 100 * (meaningful_ratio * 0.7 + (1.0 - garbage_ratio) * 0.2 + (1.0 - repeated_ratio) * 0.1)
    score = max(0, min(100, int(score)))

    issues = []
    if garbage_ratio > 0.05:
        issues.append(f"存在 {garbage_chars} 个替换字符(�)，占比 {garbage_ratio*100:.1f}%")
    if meaningful_ratio < 0.3:
        issues.append(f"有效字符占比仅 {meaningful_ratio*100:.1f}%，疑似乱码")
    if repeated_ratio > 0.1:
        issues.append(f"重复字符占比 {repeated_ratio*100:.1f}%，可能存在编码损坏")
    if len(text) < 200:
        issues.append(f"文本仅 {len(text)} 字符，内容可能不完整")
    if not issues:
        issues.append("文本质量正常")

    return {
        "score": score,
        "total_chars": total,
        "meaningful_chars": meaningful,
        "issues": issues,
    }


# ─── Adaptive Chunking: Document Profiling ────────────────────────

def profile_document(text: str) -> dict:
    """Analyze document structure and return profiling info.

    Returns:
        {
            "doc_type": "gb_standard" | "regulation" | "generic",
            "headings": [(level: int, title: str, pos: int), ...],
            "confidence": float  # 0.0 ~ 1.0
        }
    """
    lines = text.split("\n")

    # ── GB Standard detection ──
    # Pattern: numbered headings like "## 4 总则", "## 5.1 xxx", "5.3.2 xxx"
    # Also: appendix headings "## 附录A", "## A.1 xxx"
    re_gb_md_heading = re.compile(r'^(#{1,4})\s+(\d+(?:\.\d+)*)\s+(.*)$')        # ## 4 总则
    re_gb_raw_heading = re.compile(r'^(\d+(?:\.\d+)*)\s*(.{1,60})$')             # 4.1 总则 or 1范围 (no ##)
    re_gb_appendix_md = re.compile(r'^(#{1,4})\s+(附录[A-Z])\s*(.*)$')           # ## 附录A
    re_gb_appendix_sub = re.compile(r'^(#{1,4})\s+([A-Z]\.\d+)\s*(.*)$')         # ## A.1 xxx
    re_gb_raw_appendix = re.compile(r'^([A-Z]\.\d+)\s+(.{1,60})$')               # A.1 xxx (no ##)

    gb_headings = []
    for i, line in enumerate(lines):
        line_stripped = line.rstrip()
        pos = sum(len(lines[j]) + 1 for j in range(i))  # byte offset in text

        # Markdown numbered headings
        m = re_gb_md_heading.match(line_stripped)
        if m:
            level = len(m.group(1))  # number of # characters
            title = f"{m.group(2)} {m.group(3)}".strip()
            gb_headings.append((level, title, pos))
            continue

        # Appendix markdown headings: ## 附录A
        m = re_gb_appendix_md.match(line_stripped)
        if m:
            level = len(m.group(1))
            title = f"{m.group(2)} {m.group(3)}".strip() if m.group(3) else m.group(2)
            gb_headings.append((level, title, pos))
            continue

        # Appendix sub-headings markdown: ## A.1 xxx
        m = re_gb_appendix_sub.match(line_stripped)
        if m:
            level = len(m.group(1))
            title = f"{m.group(2)} {m.group(3)}".strip() if m.group(3) else m.group(2)
            gb_headings.append((level, title, pos))
            continue

        # Raw numbered headings (no markdown prefix)
        m = re_gb_raw_heading.match(line_stripped)
        if m:
            title = f"{m.group(1)} {m.group(2)}".strip()
            gb_headings.append((1, title, pos))
            continue

        # Raw appendix sub-headings (no markdown prefix): A.1 xxx
        m = re_gb_raw_appendix.match(line_stripped)
        if m:
            title = f"{m.group(1)} {m.group(2)}".strip()
            gb_headings.append((1, title, pos))
            continue

    # ── Regulation detection ──
    re_article_cn = re.compile(r'^第[一二三四五六七八九十百千零\d]+条')
    re_article_num = re.compile(r'^第(\d+)条')
    regulation_count = 0
    regulation_headings = []
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        pos = sum(len(lines[j]) + 1 for j in range(i))
        if re_article_cn.match(line_stripped) or re_article_num.match(line_stripped):
            regulation_count += 1
            # Extract article title (first 80 chars of the line)
            title = line_stripped[:80]
            regulation_headings.append((1, title, pos))

    # ── Classification ──
    # Count unique numbered sections (top-level like "1", "2", ... not sub-levels)
    top_level_gb = len(set(h[1].split()[0].split('.')[0] for h in gb_headings if h[1]))
    total_gb = len(gb_headings)

    doc_type = "generic"
    headings = []
    confidence = 0.0

    if total_gb >= 3:
        doc_type = "gb_standard"
        headings = gb_headings
        confidence = min(1.0, total_gb / 10)
    elif regulation_count >= 3:
        doc_type = "regulation"
        headings = regulation_headings
        confidence = min(1.0, regulation_count / 10)

    return {
        "doc_type": doc_type,
        "headings": headings,
        "confidence": confidence,
    }


def heading_chunk(text: str, profile: dict, min_child_size: int = 200, max_parent_size: int = 3000) -> list:
    """Split document by semantic headings. Returns same format as parent_child_chunk.

    For gb_standard:
      - child = content under each leaf heading (X.X.X or X.X level)
      - parent = content under parent heading (X level), combining all its children
      - If a child is too small (< min_child_size), merge with next sibling
      - If a parent is too large (> max_parent_size), keep it as-is (don't split further)

    For regulation:
      - child = each article (第N条)
      - parent = group of 3-5 consecutive articles

    Returns same format as parent_child_chunk:
    [{"child": str, "parent": str, "child_index": int, "parent_index": int, "section_hint": str}]
    """
    headings = profile.get("headings", [])
    doc_type = profile.get("doc_type", "generic")

    if not headings:
        return []  # caller should fall back to parent_child_chunk

    if doc_type == "gb_standard":
        return _heading_chunk_gb(text, headings, min_child_size, max_parent_size)
    elif doc_type == "regulation":
        return _heading_chunk_regulation(text, headings)
    else:
        return []


def _parse_section_number(title: str):
    """Extract the numeric part from a heading title for level comparison.

    Examples:
        "4 总则" → (4,)
        "5.1 xxx" → (5, 1)
        "5.1.2 xxx" → (5, 1, 2)
        "附录A xxx" → None (appendix)
        "A.1 xxx" → None (appendix sub)
    Returns tuple of ints or None.
    """
    m = re.match(r'(\d+(?:\.\d+)*)', title.strip())
    if m:
        return tuple(int(x) for x in m.group(1).split('.'))
    return None


def _heading_chunk_gb(text: str, headings: list, min_child_size: int = 200, max_parent_size: int = 3000) -> list:
    """Heading-based chunking for GB standard documents."""
    lines = text.split("\n")

    # Compute byte offsets for each heading position
    # (profile_document stores char offsets; we need to map to line-based splitting)
    # Re-compute line-based positions from headings
    # headings = [(level, title, pos), ...] where pos is approximate char offset

    # Build sections: for each heading, find its line index
    # We'll re-parse the text to find exact line indices of headings
    re_numbered = re.compile(r'^(#{1,4})\s+(\d+(?:\.\d+)*)\s+(.*)$')
    re_appendix_md = re.compile(r'^(#{1,4})\s+(附录[A-Z])\s*(.*)$')
    re_appendix_sub = re.compile(r'^(#{1,4})\s+([A-Z]\.\d+)\s*(.*)$')
    re_raw_numbered = re.compile(r'^(\d+(?:\.\d+)*)\s*(.{1,60})$')             # 4.1 总则 or 1范围 (no ##)
    re_raw_appendix = re.compile(r'^([A-Z]\.\d+)\s+(.{1,60})$')

    # Map each heading from profile to its line index
    heading_lines = []  # [(line_idx, level, section_number_tuple, title)]
    for level, title, _pos in headings:
        # Find the matching line
        for i, line in enumerate(lines):
            line_stripped = line.rstrip()
            matched = False
            sec_num = None

            m = re_numbered.match(line_stripped)
            if m:
                candidate_title = f"{m.group(2)} {m.group(3)}".strip()
                if candidate_title == title or title.startswith(m.group(2)):
                    sec_num = tuple(int(x) for x in m.group(2).split('.'))
                    matched = True

            if not matched:
                m = re_appendix_md.match(line_stripped)
                if m:
                    candidate_title = f"{m.group(2)} {m.group(3)}".strip() if m.group(3) else m.group(2)
                    if candidate_title == title or title.startswith(m.group(2)):
                        matched = True

            if not matched:
                m = re_appendix_sub.match(line_stripped)
                if m:
                    candidate_title = f"{m.group(2)} {m.group(3)}".strip() if m.group(3) else m.group(2)
                    if candidate_title == title or title.startswith(m.group(2)):
                        matched = True

            if not matched:
                m = re_raw_numbered.match(line_stripped)
                if m:
                    candidate_title = f"{m.group(1)} {m.group(2)}".strip()
                    if candidate_title == title or title.startswith(m.group(1)):
                        sec_num = tuple(int(x) for x in m.group(1).split('.'))
                        matched = True

            if not matched:
                m = re_raw_appendix.match(line_stripped)
                if m:
                    candidate_title = f"{m.group(1)} {m.group(2)}".strip()
                    if candidate_title == title or title.startswith(m.group(1)):
                        matched = True

            if matched:
                if sec_num is None:
                    sec_num = _parse_section_number(title)
                heading_lines.append((i, level, sec_num, title))
                break

    # Sort by line index
    heading_lines.sort(key=lambda x: x[0])

    # Post-process: assign synthetic sec_nums to appendix headings
    # so they group correctly (e.g., 附录A → (1000,), A.1 → (1000, 1), 附录B → (1001,))
    appendix_counter = [0]
    last_appendix_id = [None]
    re_appendix_title = re.compile(r'^附录([A-Z])')
    re_appendix_sub_title = re.compile(r'^([A-Z])\.(\d+)')
    updated_lines = []
    for line_idx, level, sec_num, title in heading_lines:
        if sec_num is None:
            m = re_appendix_title.match(title)
            if m:
                appendix_counter[0] += 1
                last_appendix_id[0] = appendix_counter[0]
                sec_num = (1000 + appendix_counter[0],)
            else:
                m = re_appendix_sub_title.match(title)
                if m and last_appendix_id[0] is not None:
                    sub_num = int(m.group(2))
                    sec_num = (1000 + last_appendix_id[0], sub_num)
        updated_lines.append((line_idx, level, sec_num, title))
    heading_lines = updated_lines

    if not heading_lines:
        return []

    # Split text into sections based on heading positions
    sections = []  # [(line_start, line_end, level, sec_num, title)]
    # Text before first heading
    if heading_lines[0][0] > 0:
        sections.append((0, heading_lines[0][0], 0, None, "前言"))

    for idx, (line_idx, level, sec_num, title) in enumerate(heading_lines):
        end_line = heading_lines[idx + 1][0] if idx + 1 < len(heading_lines) else len(lines)
        sections.append((line_idx, end_line, level, sec_num, title))

    # Extract text for each section
    section_texts = []
    for line_start, line_end, level, sec_num, title in sections:
        section_text = "\n".join(lines[line_start:line_end]).strip()
        section_texts.append({
            "text": section_text,
            "level": level,
            "sec_num": sec_num,
            "title": title,
            "line_start": line_start,
        })

    if not section_texts:
        return []

    # Determine leaf (child) and parent sections
    # Strategy: sections with deeper sec_num (more dots) are children
    # Sections with shallower sec_num are parents
    # Text before first heading (sec_num=None) is a parent-only section

    # Find the maximum depth of section numbers
    all_sec_nums = [s["sec_num"] for s in section_texts if s["sec_num"] is not None]
    if not all_sec_nums:
        return []

    max_depth = max(len(sn) for sn in all_sec_nums)

    # Classify: if max_depth >= 2, deeper sections (len > 1) are children,
    # top-level sections (len == 1 or None) are parents
    # If max_depth == 1, all sections are both child and parent

    # Build parent groups: group sections by their top-level number
    parent_groups = {}  # top_level_number -> [section_indices]
    current_top = None
    for i, s in enumerate(section_texts):
        if s["sec_num"] is not None:
            top = s["sec_num"][0]
            if top != current_top:
                current_top = top
            if current_top not in parent_groups:
                parent_groups[current_top] = []
            parent_groups[current_top].append(i)
        else:
            # Text before first heading — standalone parent
            parent_groups[None] = [i]

    # Merge small children
    for top_level, indices in parent_groups.items():
        if len(indices) <= 1:
            continue
        merged = []
        for idx in indices:
            if merged and len(section_texts[idx]["text"]) < min_child_size:
                # Merge into previous
                merged[-1] = idx
            else:
                merged.append(idx)
        parent_groups[top_level] = merged

    # Build chunks
    results = []
    child_index = 0
    parent_index = 0

    # Process sections in order
    processed_parents = set()

    for i, s in enumerate(section_texts):
        top = s["sec_num"][0] if s["sec_num"] is not None else None
        if top in processed_parents:
            continue

        if top is None:
            # Pre-heading text: standalone parent
            parent_text = s["text"]
            if not parent_text.strip():
                processed_parents.add(top)
                continue

            # This is a child under this parent
            section_hint = s["title"][:80] if s["title"] else parent_text[:80]
            # If child too small, merge into parent directly
            if len(parent_text) < min_child_size and len(section_texts) > 1:
                # Merge with next section's parent
                next_i = i + 1
                if next_i < len(section_texts):
                    next_parent_text = "\n\n".join(
                        section_texts[j]["text"] for j in parent_groups.get(
                            section_texts[next_i]["sec_num"][0] if section_texts[next_i]["sec_num"] else None, []
                        )
                    )
                    parent_text = parent_text + "\n\n" + next_parent_text[:max_parent_size]

            results.append({
                "child": parent_text[:800],
                "parent": parent_text[:max_parent_size],
                "child_index": child_index,
                "parent_index": parent_index,
                "section_hint": section_hint,
            })
            child_index += 1
            parent_index += 1
            processed_parents.add(top)
            continue

        # Get all sections under this parent
        group_indices = parent_groups.get(top, [])
        if not group_indices:
            continue

        # Check if this is a leaf section (no children)
        is_leaf = max_depth == 1 or len(s["sec_num"]) == max_depth

        if is_leaf and len(group_indices) == 1:
            # Single leaf section: child = this section, parent = this section
            parent_text = s["text"]
            section_hint = s["title"][:80] if s["title"] else parent_text[:80]

            results.append({
                "child": parent_text[:800],
                "parent": parent_text[:max_parent_size],
                "child_index": child_index,
                "parent_index": parent_index,
                "section_hint": section_hint,
            })
            child_index += 1
            parent_index += 1
        else:
            # Parent section with children
            all_text = "\n\n".join(section_texts[j]["text"] for j in group_indices if section_texts[j]["text"].strip())
            section_hint = s["title"][:80] if s["title"] else all_text[:80]

            # Create child chunks from individual sections
            for j in group_indices:
                child_text = section_texts[j]["text"]
                if not child_text.strip():
                    continue
                if len(child_text) < min_child_size:
                    # Try to merge with next sibling
                    next_j_idx = group_indices.index(j) + 1
                    if next_j_idx < len(group_indices):
                        next_j = group_indices[next_j_idx]
                        child_text = child_text + "\n\n" + section_texts[next_j]["text"]
                    if len(child_text) < min_child_size:
                        # Still too small, use as-is
                        pass

                results.append({
                    "child": child_text[:800],
                    "parent": all_text[:max_parent_size],
                    "child_index": child_index,
                    "parent_index": parent_index,
                    "section_hint": section_hint,
                })
                child_index += 1

            parent_index += 1

        processed_parents.add(top)

    return results


def _heading_chunk_regulation(text: str, headings: list) -> list:
    """Heading-based chunking for regulation documents (article-based)."""
    lines = text.split("\n")
    re_article = re.compile(r'^(第[一二三四五六七八九十百千零\d]+条)')

    # Find article line indices
    article_lines = []
    for level, title, _pos in headings:
        for i, line in enumerate(lines):
            if re.match(r'^' + re.escape(title[:5]).replace(r'\ ', ' ') , line.strip()):
                article_lines.append((i, title))
                break
        else:
            # Fallback: find by article number pattern
            for i, line in enumerate(lines):
                if re_article.match(line.strip()):
                    article_lines.append((i, title))
                    break

    if not article_lines:
        return []

    # Sort by line index
    article_lines.sort(key=lambda x: x[0])

    # Split into article texts
    article_texts = []
    for idx, (line_idx, title) in enumerate(article_lines):
        end_line = article_lines[idx + 1][0] if idx + 1 < len(article_lines) else len(lines)
        article_text = "\n".join(lines[line_idx:end_line]).strip()
        article_texts.append({"text": article_text, "title": title})

    if not article_texts:
        return []

    # child = each article, parent = group of 3-5 articles
    results = []
    child_index = 0
    parent_index = 0
    group_size = 4  # articles per parent group

    for i in range(0, len(article_texts), group_size):
        group = article_texts[i:i + group_size]
        parent_text = "\n\n".join(a["text"] for a in group)
        section_hint = group[0]["title"][:80]

        for a in group:
            results.append({
                "child": a["text"][:800],
                "parent": parent_text,
                "child_index": child_index,
                "parent_index": parent_index,
                "section_hint": section_hint,
            })
            child_index += 1

        parent_index += 1

    return results


def parent_child_chunk(text: str, child_size: int = 384, parent_size: int = 2048, overlap: int = 80) -> list:
    """将文本切分为父子分块。

    返回 list of dict:
    [
        {
            "child": "子块文本（用于向量匹配）",
            "parent": "父块文本（用于LLM上下文）",
            "child_index": 0,
            "parent_index": 0,
            "section_hint": "父块前80字符（章节提示）"
        },
        ...
    ]
    """
    # Step 1: 按段落分割
    paragraphs = text.split("\n\n")
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    # Step 2: 聚合段落为父块（按 parent_size 滑动窗口）
    parents = []
    p_idx = 0
    while p_idx < len(paragraphs):
        parent_text = ""
        while p_idx < len(paragraphs):
            candidate = (parent_text + "\n\n" + paragraphs[p_idx]).strip() if parent_text else paragraphs[p_idx]
            if len(candidate) > parent_size and parent_text:
                break
            parent_text = candidate
            p_idx += 1
        if parent_text:
            parents.append(parent_text)

    # Step 3: 在每个父块内切子块
    results = []
    child_index = 0
    for p_idx, parent_text in enumerate(parents):
        section_hint = parent_text[:80]
        # 按 child_size 滑动窗口切子块（子块可以跨段落边界）
        pos = 0
        while pos < len(parent_text):
            end = min(pos + child_size, len(parent_text))
            child_text = parent_text[pos:end]
            if child_text.strip():
                results.append({
                    "child": child_text,
                    "parent": parent_text,
                    "child_index": child_index,
                    "parent_index": p_idx,
                    "section_hint": section_hint,
                })
                child_index += 1
            pos += child_size - overlap  # 带重叠的滑动窗口
            if pos >= len(parent_text):
                break

    return results


# ─── API ──────────────────────────────────────────────────────
@app.post("/api/upload")
async def upload(
    file: UploadFile = File(...),
    title: str = Form(""),
    category: str = Form(""),
    bank: str = Form("kb"),
    confirm_quality: str = Form(""),
):
    """上传文档 → 解析 → 切片 → 存入 Hindsight（支持指定 bank）"""
    if not file.filename:
        raise HTTPException(400, "文件名不能为空")

    # 验证 bank
    bank_cfg = get_bank_config(bank)
    if bank == "all":
        bank = "kb"  # "全部"默认回到 kb
    hs_bank = bank_cfg["hindsight"]

    content = await file.read()
    if len(content) == 0:
        raise HTTPException(400, "文件为空")


    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(400, f"文件过大（{len(content)//1024//1024}MB），上限 {MAX_FILE_SIZE//1024//1024}MB")

    try:
        text = await parse_document(file.filename, content)
        # 清除空字节，PostgreSQL UTF8 不接受 0x00
        text = text.replace("\x00", "")
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        traceback.print_exc(file=sys.stderr)
        raise HTTPException(500, f"文档解析异常: {e}")

    if not text or len(text.strip()) < 10:
        raise HTTPException(400, "文档内容过短")

    # ── 质量评估 ──
    quality = assess_quality(text)
    if quality["score"] < 80 and confirm_quality != "true":
        return {
            "ok": False,
            "detail": f"文档解析质量较低（{quality['score']}%），可能存在乱码。建议检查后重新上传或使用 MinerU 解析。",
            "quality": {
                "score": quality["score"],
                "issues": quality["issues"],
                "needs_confirm": True,
            }
        }

    # ── 去重检测：按文件内容 SHA256 查重 ──
    content_hash = hashlib.sha256(content).hexdigest()
    existing = find_by_hash(content_hash)
    if existing:
        raise HTTPException(
            409,
            f"文档已存在（内容完全一致）。\n"
            f"已有文档：{existing['title']}（{existing['filename']}）\n"
            f"上传时间：{existing['created_at']}\n"
            f"文档ID：{existing['doc_id']}"
        )

    # 标题：用户指定 > 文件名去扩展名
    doc_title = title.strip() or filename_to_title(file.filename)
    doc_category = category.strip()

    # 分块存入 Hindsight（批量一次提交）
    # ── Adaptive chunking: profile document first ──
    profile = profile_document(text)
    doc_type = profile.get("doc_type", "generic")
    print(f"Document profile: type={doc_type}, confidence={profile.get('confidence', 0):.2f}, headings={len(profile.get('headings', []))}", file=sys.stderr)

    if doc_type == "gb_standard" and profile.get("confidence", 0) >= 0.3:
        pc_chunks = heading_chunk(text, profile)
        if pc_chunks:
            print(f"Heading-based chunking: {len(pc_chunks)} chunks", file=sys.stderr)
        else:
            # Fallback: heading parsing produced no chunks
            doc_type = "generic"
            pc_chunks = parent_child_chunk(text, child_size=384, parent_size=2048, overlap=80)
            print(f"Heading chunking returned 0, fell back to paragraph: {len(pc_chunks)} chunks", file=sys.stderr)
    elif doc_type == "regulation" and profile.get("confidence", 0) >= 0.3:
        pc_chunks = heading_chunk(text, profile)
        if pc_chunks:
            print(f"Regulation-based chunking: {len(pc_chunks)} chunks", file=sys.stderr)
        else:
            doc_type = "generic"
            pc_chunks = parent_child_chunk(text, child_size=384, parent_size=2048, overlap=80)
            print(f"Regulation chunking returned 0, fell back to paragraph: {len(pc_chunks)} chunks", file=sys.stderr)
    else:
        pc_chunks = parent_child_chunk(text, child_size=384, parent_size=2048, overlap=80)
        print(f"Paragraph-based chunking: {len(pc_chunks)} chunks", file=sys.stderr)

    doc_id = str(uuid.uuid4())

    # 构建 parent 映射（用于查询时检索 parent 上下文）
    parent_map = {}  # parent_index -> parent_text
    for pc in pc_chunks:
        parent_map[pc["parent_index"]] = pc["parent"]

    memory_items = []
    for i, pc in enumerate(pc_chunks):
        child_content = pc["child"].strip()
        if not child_content:
            continue
        # 子块内容加上 section hint 前缀提高匹配质量
        enhanced_content = f"[{pc['section_hint']}] {child_content}" if pc["section_hint"] else child_content

        tags = [
            f"doc:{file.filename}",
            f"chunk:{i+1}/{len(pc_chunks)}",
            f"doc_id:{doc_id}",
            f"title:{doc_title}",
            f"bank:{bank}",
            f"parent_idx:{pc['parent_index']}",
            f"strategy:{doc_type}",
        ]
        if doc_category:
            tags.append(f"cat:{doc_category}")
        memory_items.append({"content": enhanced_content, "tags": tags, "type": "world"})

    retained = 0
    hindsight_error = None
    if memory_items:
        # 超时按 chunk 数量动态计算：每 chunk 5s，最少 120s，最多 600s
        # 208 chunks 实测需要 ~408s，5s/chunk 给 2.5× 余量
        dyn_timeout = max(120, min(len(memory_items) * 5, 600))
        print(f"Upload: {len(memory_items)} chunks → bank={hs_bank}, timeout={dyn_timeout}s", file=sys.stderr)
        try:
            result = await hindsight_request(
                f"/v1/default/banks/{hs_bank}/memories",
                "POST",
                {"items": memory_items},
                timeout=dyn_timeout,
            )
            retained = result.get("items_count", len(memory_items))
            print(f"Hindsight stored {retained}/{len(memory_items)} chunks for doc {doc_id}", file=sys.stderr)
        except Exception as e:
            hindsight_error = str(e) or repr(e)
            print(f"Hindsight write FAILED for doc {doc_id} ({file.filename}): {hindsight_error}", file=sys.stderr)

    # 如果 Hindsight 完全失败，不存元数据，返回错误
    if hindsight_error and retained == 0:
        raise HTTPException(
            502,
            f"知识库存储服务暂时不可用，请稍后重试。"
            f"（已将 {len(memory_items)} 个文本块上传到向量库但全部失败）"
        )

    # 部分成功：存元数据，返回警告
    if retained < len(memory_items):
        print(f"WARNING: Only {retained}/{len(memory_items)} chunks stored for doc {doc_id}", file=sys.stderr)

    # 保存元数据到 SQLite（含 bank 字段）
    save_meta(doc_id, doc_title, doc_category, file.filename, content_hash, doc_type)

    # 保存 parent_map 到 SQLite（供查询时检索父块上下文）
    try:
        db = get_db()
        for idx, ptext in parent_map.items():
            db.execute("INSERT OR REPLACE INTO parent_chunks (doc_id, parent_idx, parent_text) VALUES (?, ?, ?)",
                       (doc_id, idx, ptext))
        db.commit()
        db.close()
        print(f"Saved {len(parent_map)} parent chunks for doc {doc_id}", file=sys.stderr)
    except Exception as e:
        print(f"WARNING: Failed to save parent_chunks for {doc_id}: {e}", file=sys.stderr)

    # 保存PDF原件副本到备份目录（去重检查通过后才备份）
    backup_dir = os.path.join(BASE_DIR, "storage", "backups")
    os.makedirs(backup_dir, exist_ok=True)
    backup_name = file.filename or "unknown.pdf"
    backup_path = os.path.join(backup_dir, backup_name)
    try:
        with open(backup_path, "wb") as bf:
            bf.write(content)
        print(f"Backup saved: {backup_path} ({len(content)//1024}KB)", file=sys.stderr)
    except Exception as e:
        print(f"WARNING: PDF backup failed: {e}", file=sys.stderr)

    # 同时记录 bank 到 meta.db
    db = get_db()
    db.execute("UPDATE doc_meta SET bank = ? WHERE doc_id = ?", (bank, doc_id))
    db.commit()
    db.close()

    # ── 完整性验证：上传后召回比对 ──
    integrity = None
    try:
        await asyncio.sleep(2)  # 等待 Hindsight 索引
        recalled_chunks = await recall(doc_title, limit=50, bank=hs_bank, max_tokens=32768)
        recalled_text = "\n".join(r.get("text", "") for r in recalled_chunks)
        if recalled_text and len(recalled_text) > 200:
            coverage = min(100, round(len(recalled_text) / max(len(text), 1) * 100, 1))
            integrity = {
                "original_chars": len(text),
                "recalled_chars": len(recalled_text),
                "coverage_pct": coverage,
                "status": "ok" if coverage >= 80 else ("partial" if coverage >= 50 else "low"),
            }
        else:
            integrity = {
                "original_chars": len(text),
                "recalled_chars": len(recalled_text) if recalled_text else 0,
                "coverage_pct": 0,
                "status": "pending",
                "note": "索引尚未完成，请稍后查看",
            }
    except Exception as e:
        print(f"Upload integrity check failed for {doc_id}: {e}", file=sys.stderr)

    return {
        "ok": True,
        "doc_id": doc_id,
        "title": doc_title,
        "category": doc_category,
        "filename": file.filename,
        "chunks": retained,
        "total_chars": len(text),
        "preview": text[:200] + ("..." if len(text) > 200 else ""),
        "warning": f"仅成功入库 {retained}/{len(memory_items)} 个文本片段" if retained < len(memory_items) else None,
        "quality": {
            "score": quality["score"],
            "issues": quality["issues"],
            "needs_confirm": quality["score"] < 80,
        },
        "integrity": integrity,
        "doc_type": doc_type,
    }


@app.post("/api/query")
async def query(q: str = Form(...), bank: str = Form("all"), history: str = Form("")):
    """搜索知识库 → 召回 → DeepSeek 合成答案（支持多 bank）"""
    if not q.strip():
        raise HTTPException(400, "问题不能为空")

    bank_cfg = get_bank_config(bank)
    bank_prompt = bank_cfg["prompt"]

    # 构建 doc_id → bank 映射（用于过滤结果）
    db = get_db()
    bank_map = {}
    try:
        rows = db.execute("SELECT doc_id, bank FROM doc_meta").fetchall()
        bank_map = {r["doc_id"]: r["bank"] for r in rows}
    except sqlite3.OperationalError:
        pass

    # ── 精确检索：关键词匹配 meta.db 标题 ──
    import re
    exact_results = []
    # 提取查询中的标准号/文件号等精确标识符
    patterns = [
        r'GB/T\s*\d+[\.\-]\d+', r'GB\s*\d+[\.\-]\d+',
        r'T/EGAG\s*\d+[\.\-]\d+', r'GDZW\s*\d+[\.\-]\d+',
        r'粤府办〔\d+〕\d+号', r'穗政数〔\d+〕\d+号',
        r'ISO\s*\d+', r'[\u4e00-\u9fff]+〔\d+〕\d+号',
    ]
    exact_terms = set()
    for pat in patterns:
        exact_terms.update(re.findall(pat, q))
    
    if exact_terms:
        # 查 meta.db 标题匹配
        title_rows = db.execute(
            "SELECT doc_id, title FROM doc_meta WHERE " + " OR ".join(["title LIKE ?" for _ in exact_terms]),
            [f"%{t}%" for t in exact_terms]
        ).fetchall()
        db.close()
        
        for tr in title_rows:
            # 过滤 bank
            if bank != "all" and bank_map.get(tr["doc_id"]) != bank:
                continue
            if bank_map.get(tr["doc_id"]) == "skip":
                continue
            # 用文档标题做定向召回
            targeted = await recall(tr["title"], limit=2, bank="kb")
            for r in targeted:
                tags = r.get("tags", [])
                doc_tag = None
                for t in tags:
                    if t.startswith("doc_id:"):
                        doc_tag = t[7:]
                        break
                if doc_tag == tr["doc_id"]:
                    exact_results.append(r)
    else:
        db.close()

    # ── Hybrid Search: Dense + BM25 RRF 融合 ──
    raw_results = await recall(q, limit=25, bank="kb")
    # "全部"模式增加召回深度
    if bank == "all":
        try:
            extra = await recall(q, limit=15, bank="kb")
            raw_results = extra + raw_results
        except Exception:
            pass

    # BM25 关键词召回（补充语义召回的盲区）
    bm25_merged = raw_results  # 默认 fallback
    try:
        bm25_index, bm25_docs = await build_bm25_index(bank)
        if bm25_index:
            bm25_hits = bm25_search(q, bm25_index, bm25_docs, top_k=20)
            if bm25_hits:
                bm25_merged = rrf_merge(raw_results, bm25_hits, k=60)
    except Exception as e:
        print(f"[WARN] BM25 fallback: {e}", flush=True)

    # 精确结果排在最前面
    all_results = exact_results + bm25_merged

    # ── 清洗 + 过滤 + 去重合并 ──
    doc_facts = {}  # doc_id → [(text, doc_name, cleaned_text), ...]
    
    for r in all_results:
        text = r.get("text", "") or ""
        tags = r.get("tags", [])
        
        # 提取 doc_id
        doc_id = None
        for t in tags:
            if t.startswith("doc_id:"):
                doc_id = t[7:]
                break
        if not doc_id:
            # Hindsight consolidation 后所有 bank 的记忆都可能丢失 doc_id tag，不能丢弃
            doc_id = f"_notag_{id(r)}"  # 伪造唯一 ID 用于分组
        # 过滤 skip bank
        if bank_map.get(doc_id) == "skip":
            continue
        # 过滤指定 bank
        if bank != "all" and bank_map.get(doc_id) and bank_map.get(doc_id) != bank:
            continue

        # 提取文档名
        doc_name = "未知文档"
        for t in tags:
            if t.startswith("title:"):
                doc_name = t[6:]
                break
        
        # 清理 Hindsight 元数据（| When: ... | Involving: ...）
        cleaned = re.sub(r'\s*\|\s*(When|Involving|Entities|Location|Type|Source):[^|]*', '', text).strip()
        if not cleaned:
            cleaned = text.strip()
        
        # 提取 parent_idx（parent-child 分块机制）
        parent_idx = None
        for t in tags:
            if t.startswith("parent_idx:"):
                try:
                    parent_idx = int(t.split(":", 1)[1])
                except (ValueError, IndexError):
                    pass
                break
        if doc_id not in doc_facts:
            doc_facts[doc_id] = []
        doc_facts[doc_id].append((text, doc_name, cleaned, parent_idx))

    if not doc_facts:
        return {"answer": "知识库中未找到相关信息。", "sources": []}

    # ── 按文档合并：每个文档取 top-2 fact，拼接 ──
    context_parts = []
    sources = []

    # ── 批量查询 parent 上下文 ──
    parent_text_cache = {}  # (doc_id, parent_idx) -> parent_text
    parent_keys_to_fetch = set()
    for doc_id, facts in doc_facts.items():
        for _, _, _, pidx in facts[:3]:
            if pidx is not None:
                parent_keys_to_fetch.add((doc_id, pidx))
    if parent_keys_to_fetch:
        try:
            pdb = get_db()
            for did, pidx in parent_keys_to_fetch:
                row = pdb.execute("SELECT parent_text FROM parent_chunks WHERE doc_id=? AND parent_idx=?",
                                  (did, pidx)).fetchone()
                if row:
                    parent_text_cache[(did, pidx)] = row[0]
            pdb.close()
        except Exception as e:
            print(f"[WARN] parent_chunks query failed: {e}", file=sys.stderr)
    
    for doc_id, facts in doc_facts.items():
        # 取前 2 个 fact
        top_facts = facts[:3]
        doc_name = top_facts[0][1]

        # 尝试用 parent 上下文（更完整）替代 child（更碎片化）
        parent_texts_for_doc = []
        seen_parent = set()
        for _, _, _, pidx in top_facts:
            if pidx is not None and (doc_id, pidx) in parent_text_cache:
                pt = parent_text_cache[(doc_id, pidx)]
                if pt[:80] not in seen_parent:
                    seen_parent.add(pt[:80])
                    parent_texts_for_doc.append(pt)

        if parent_texts_for_doc:
            combined = "\n\n".join(parent_texts_for_doc[:3])
        else:
            # Fallback: 用 child 文本（去重合并）
            seen_texts = set()
            merged = []
            for _, _, cleaned, _ in top_facts:
                key = cleaned[:80]
                if key not in seen_texts:
                    seen_texts.add(key)
                    merged.append(cleaned)
            if not merged:
                continue
            combined = "；".join(merged)

        context_parts.append(f"[来源: {doc_name}]\n{combined}")

        # Merge all top facts' cleaned text (up to 800 chars) for sources display
        merged_text = "；".join([c for _, _, c, _ in facts[:3]])
        sources.append({
            "doc": doc_name,
            "doc_id": doc_id if not doc_id.startswith("_notag_") else None,
            "chunk": f"{len(facts)} 条相关",
            "text": merged_text[:800],
        })
    
    # ── 限制 context 总量 ──
    total_chars = sum(len(p) for p in context_parts)
    if total_chars > 10000:
        # 按序截断
        kept = []
        chars = 0
        for p in context_parts:
            if chars + len(p) > 8000:
                break
            kept.append(p)
            chars += len(p)
        context_parts = kept
    
    context = "\n\n---\n\n".join(context_parts)
    sources = sources[:12]

    # ── 追问上下文注入 ──
    history_context = ""
    if history.strip():
        try:
            history_context = f"\n\n【对话历史】\n{history.strip()}\n\n请结合上述对话历史理解当前问题。如果当前问题是追问，基于历史上下文中断的地方继续回答。"
        except Exception:
            pass
    
    prompt = f"""{bank_prompt}

【回答原则】
1. 以「文档内容」为主要依据，优先引用文档中的具体内容和数据
2. 可以基于文档内容进行综合推理和归纳总结，但不得编造文档中不存在的具体数字、条款号或标准编号
3. 每个关键论断标注来源文档名称
4. 如果文档中完全没有相关信息（关键词完全不匹配），说明"当前知识库未收录相关内容，以下为一般性参考"，然后基于你的知识给出方向性建议
5. 多个文档存在矛盾时，列出不同说法并各自标注来源

基于以下文档内容回答问题：

文档内容：
{context}
{history_context}

问题：{q}

请用中文回答，引用具体条款和数据，并标注信息来源。"""

    try:
        answer = await deepseek_chat([
            {"role": "system", "content": bank_prompt},
            {"role": "user", "content": prompt},
        ])
    except Exception as e:
        answer = f"答案生成失败: {e}"

    return {"answer": answer, "sources": sources}


async def web_search(query: str, max_results: int = 3) -> str:
    """通过 AnySearch CLI 联网搜索，返回格式化的搜索结果文本"""
    try:
        skill_dir = os.path.expanduser("~/.agents/skills/anysearch")
        cli_path = os.path.join(skill_dir, "scripts", "anysearch_cli.py")
        if not os.path.exists(cli_path):
            return ""

        proc = await asyncio.create_subprocess_exec(
            "python3", cli_path, "search", query, "--max_results", str(max_results),
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=15)
        if proc.returncode == 0 and stdout:
            return stdout.decode("utf-8", errors="replace").strip()
        return ""
    except Exception as e:
        print(f"[WARN] web_search failed: {e}", flush=True)
        return ""


@app.post("/api/web-search")
async def web_search_api(
    q: str = Form(...),
    bank: str = Form("all"),
    context: str = Form(""),
):
    """联网搜索 — 用户对知识库结果不满意时，结合页面上下文联网搜索回答"""
    if not q.strip():
        raise HTTPException(400, "问题不能为空")

    bank_cfg = get_bank_config(bank)
    bank_prompt = bank_cfg["prompt"]

    # 联网搜索
    web_results = await web_search(q, max_results=5)
    web_context = ""
    if web_results:
        web_context = f"\n\n【联网搜索结果】\n{web_results}\n"

    prompt = f"""{bank_prompt}

用户在一个知识库问答系统中搜索了以下问题，但对知识库返回的结果不满意，需要联网搜索补充。

【用户原始问题】
{q}

【当前页面已有的知识库答案】
{context[:3000] if context else '(无)'}
{web_context}
请综合以上信息，优先参考联网搜索结果，结合知识库已有内容，给出完整、准确的回答。
- 引用具体条款和数据
- 标注信息来源
- 如果联网搜索也没有找到，基于你的知识给出方向性建议
直接回答，用中文。"""

    try:
        answer = await deepseek_chat([
            {"role": "system", "content": bank_prompt},
            {"role": "user", "content": prompt},
        ])
    except Exception as e:
        answer = f"联网搜索回答失败: {e}"

    return {"answer": answer, "web_searched": bool(web_results)}


@app.get("/api/documents/{doc_id}/content")
async def get_document_content(doc_id: str):
    """返回文档的完整文本内容"""
    # 从 meta.db 查找文档所属的 bank
    db = get_db()
    meta = db.execute("SELECT title, filename, bank, created_at FROM doc_meta WHERE doc_id = ?", (doc_id,)).fetchone()
    db.close()
    
    if not meta:
        raise HTTPException(404, "文档不存在")
    
    doc_bank = meta["bank"] or "kb"
    bank_cfg = get_bank_config(doc_bank)
    hs_bank = bank_cfg["hindsight"]  # 映射到 Hindsight bank
    
    # 尝试主 Hindsight bank
    docs_result = await hindsight_request(
        f"/v1/default/banks/{hs_bank}/documents",
        timeout=10
    )
    doc_list = docs_result.get("items", []) or docs_result.get("documents", [])
    
    hindsight_doc_id = None
    for d in doc_list:
        tags = d.get("tags", [])
        if f"doc_id:{doc_id}" in tags:
            hindsight_doc_id = d.get("id")
            break
    
    # 主 bank 未找到 → 回退旧 Hindsight bank 名
    OLD_HINDSIGHT_BANKS = ["tech", "security", "ai", "notes", "proposals", "assessment", "projects"]
    if not hindsight_doc_id:
        for fallback in OLD_HINDSIGHT_BANKS:
            if fallback == hs_bank:
                continue
            try:
                fb_result = await hindsight_request(
                    f"/v1/default/banks/{fallback}/documents",
                    timeout=5
                )
                fb_list = fb_result.get("items", []) or fb_result.get("documents", [])
                for d in fb_list:
                    if f"doc_id:{doc_id}" in d.get("tags", []):
                        hindsight_doc_id = d.get("id")
                        hs_bank = fallback  # 用找到的 bank
                        break
            except Exception:
                pass
            if hindsight_doc_id:
                break
    
    if not hindsight_doc_id:
        # fallback: semantic recall by title
        try:
            title = meta["title"] or ""
            if title:
                recalled = await recall(title, limit=50, bank="kb", max_tokens=32768)
                if recalled:
                    full_text = "\n\n".join(r.get("text", "") for r in recalled)
                    if full_text and len(full_text) > 50:
                        return {
                            "doc_id": doc_id,
                            "title": title,
                            "filename": meta["filename"] or "",
                            "chunks": len(recalled),
                            "text": full_text,
                            "source": "recall",  # marker: content from semantic recall, not original
                        }
        except Exception:
            pass
        raise HTTPException(404, "文档内容未找到（可能尚未完成索引）")
    
    # 获取单个文档的完整内容（含 original_text，用正确的 hs_bank）
    doc_detail = await hindsight_request(
        f"/v1/default/banks/{hs_bank}/documents/{hindsight_doc_id}",
        timeout=10
    )
    full_text = doc_detail.get("original_text", "") or doc_detail.get("text", "") or ""
    chunks_count = doc_detail.get("memory_unit_count", 0)
    
    return {
        "doc_id": doc_id,
        "title": meta["title"] if meta else "未知文档",
        "filename": meta["filename"] if meta else "",
        "chunks": chunks_count,
        "text": full_text,
    }


@app.post("/api/fetch-standard")
async def fetch_standard(std_no: str = Form(...), bank: str = Form("kb")):
    """从公开来源下载国家标准并入库"""
    import subprocess, tempfile, os as _os_module
    
    if not std_no.strip():
        raise HTTPException(400, "标准号不能为空")
    
    bank_cfg = get_bank_config(bank)
    if bank == "all":
        bank = "kb"
    hs_bank = bank_cfg["hindsight"]
    
    # Step 1: Search with AnySearch
    anysearch_cli = _os_module.path.expanduser("~/.agents/skills/anysearch/scripts/anysearch_cli.py")
    try:
        result = subprocess.run(
            ["python3", anysearch_cli, "search", f"{std_no} 标准 PDF下载", "-m", "5", "--freshness", "year"],
            capture_output=True, text=True, timeout=30
        )
    except subprocess.TimeoutExpired:
        raise HTTPException(500, "搜索超时")
    except Exception as e:
        raise HTTPException(500, f"搜索失败: {e}")
    
    # Parse URLs from markdown output
    urls = [line.split("**URL**: ")[1].strip() for line in result.stdout.split('\n') if "**URL**: " in line]
    if not urls:
        raise HTTPException(404, f"未找到 {std_no} 的下载链接")
    
    # Step 2: Download PDF
    pdf_path = None
    import httpx as _httpx_module
    async with _httpx_module.AsyncClient(timeout=60, follow_redirects=True) as client:
        for url in urls[:3]:
            try:
                resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
                if resp.status_code == 200 and "application/pdf" in resp.headers.get("content-type", ""):
                    fd, pdf_path = tempfile.mkstemp(suffix=".pdf")
                    _os_module.write(fd, resp.content)
                    _os_module.close(fd)
                    break
            except Exception:
                continue
    
    if not pdf_path:
        raise HTTPException(500, "下载失败：所有链接均不可用")
    
    # Step 3: Parse PDF
    try:
        text = ""
        with open(pdf_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        _os_module.unlink(pdf_path)
    except Exception as e:
        if _os_module.path.exists(pdf_path):
            _os_module.unlink(pdf_path)
        raise HTTPException(500, f"PDF解析失败: {e}")
    
    if not text or len(text.strip()) < 100:
        raise HTTPException(400, "PDF内容过短，可能是扫描件")
    
    # Step 4: Upload to Hindsight
    doc_title = std_no.strip()
    chunk_size = 1000
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    doc_id = str(uuid.uuid4())
    
    memory_items = []
    for i, chunk in enumerate(chunks):
        chunk = chunk.strip()
        if not chunk:
            continue
        tags = [
            f"doc:{doc_title}.pdf",
            f"chunk:{i+1}/{len(chunks)}",
            f"doc_id:{doc_id}",
            f"title:{doc_title}",
            f"bank:{bank}",
        ]
        memory_items.append({"content": chunk, "tags": tags, "type": "world"})
    
    success_count = 0
    total_count = len(memory_items)
    if memory_items:
        dyn_timeout = max(120, min(len(memory_items) * 5, 600))
        try:
            result = await hindsight_request(
                f"/v1/default/banks/{hs_bank}/memories",
                "POST",
                {"items": memory_items},
                timeout=dyn_timeout,
            )
            success_count = result.get("items_count", 0)
        except Exception as e:
            raise HTTPException(500, f"入库失败: {e}")
    
    if success_count == 0:
        raise HTTPException(500, "入库失败：所有chunk写入Hindsight均失败")
    
    # Step 5: Write meta.db
    db = get_db()
    db.execute(
        "INSERT OR REPLACE INTO doc_meta (doc_id, title, category, filename, content_hash, bank, created_at) VALUES (?, ?, '', ?, '', ?, datetime('now'))",
        (doc_id, doc_title, f"{doc_title}.pdf", bank)
    )
    db.commit()
    db.close()
    
    return {
        "ok": True,
        "doc_id": doc_id,
        "title": doc_title,
        "bank": bank,
        "text_length": len(text),
        "chunks": f"{success_count}/{total_count}",
    }


@app.get("/api/documents")
async def list_documents(bank: str = "all"):
    """列出文档（meta.db 为主，Hindsight 补充 chunk/size）"""
    db = get_db()
    
    # 从 meta.db 查文档（源数据）
    if bank == "all":
        meta_rows = db.execute(
            "SELECT doc_id, title, category, filename, bank, created_at FROM doc_meta WHERE bank NOT IN ('skip') ORDER BY created_at DESC"
        ).fetchall()
    else:
        meta_rows = db.execute(
            "SELECT doc_id, title, category, filename, bank, created_at FROM doc_meta WHERE bank = ? ORDER BY created_at DESC",
            (bank,)
        ).fetchall()
    db.close()
    
    # 从 Hindsight 补充 chunk/size 数据
    result = await hindsight_request(f"/v1/default/banks/kb/documents?limit=1000", timeout=15)
    hs_stats = {}  # doc_id → {chunks, size}
    for item in result.get("items", []):
        doc_id = None
        for t in item.get("tags", []):
            if t.startswith("doc_id:"):
                doc_id = t[7:]
                break
        if not doc_id:
            continue
        if doc_id not in hs_stats:
            hs_stats[doc_id] = {"chunks": 0, "size": 0}
        hs_stats[doc_id]["chunks"] += 1
        hs_stats[doc_id]["size"] += item.get("text_length", 0)
    
    docs = []
    for r in meta_rows:
        stats = hs_stats.get(r["doc_id"], {"chunks": 0, "size": 0})
        docs.append({
            "id": r["doc_id"],
            "title": r["title"] or "未知文档",
            "category": r["category"] or "",
            "filename": r["filename"] or "",
            "chunks": stats["chunks"],
            "size_chars": stats["size"],
            "created": r["created_at"] or "",
            "bank": r["bank"] or "kb",
        })
    
    return {"documents": docs}


@app.patch("/api/documents/{doc_id}")
async def patch_document(doc_id: str, title: str = Form(None), category: str = Form(None)):
    """编辑文档的标题和分类"""
    if not title and not category:
        raise HTTPException(400, "至少需要提供 title 或 category")
    update_meta(doc_id, title=title, category=category)
    return {"ok": True, "doc_id": doc_id, "title": title, "category": category}


@app.patch("/api/documents/{doc_id}/bank")
async def patch_document_bank(doc_id: str, bank: str = Form(...)):
    """切换文档所属 bank"""
    if bank not in BANKS:
        raise HTTPException(400, f"无效的 bank: {bank}，有效值: {', '.join(BANKS.keys())}")
    db = get_db()
    row = db.execute("SELECT doc_id FROM doc_meta WHERE doc_id = ?", (doc_id,)).fetchone()
    if not row:
        db.close()
        raise HTTPException(404, f"文档 {doc_id} 不存在")
    db.execute("UPDATE doc_meta SET bank = ? WHERE doc_id = ?", (bank, doc_id))
    db.commit()
    db.close()
    return {"ok": True, "doc_id": doc_id, "bank": bank}


@app.get("/api/wiki")
async def wiki_tree(bank: str = "all"):
    """返回知识库目录树结构，按 bank → category → documents 组织"""
    db = get_db()
    
    if bank == "all":
        rows = db.execute(
            "SELECT doc_id, title, category, filename, bank, created_at, doc_type FROM doc_meta WHERE bank NOT IN ('skip') ORDER BY bank, category, title"
        ).fetchall()
    else:
        rows = db.execute(
            "SELECT doc_id, title, category, filename, bank, created_at, doc_type FROM doc_meta WHERE bank = ? AND bank NOT IN ('skip') ORDER BY category, title",
            (bank,)
        ).fetchall()
    db.close()
    
    # Build tree: bank → category → [documents]
    tree = {}
    bank_counts = {}
    for r in rows:
        b = r["bank"] or "kb"
        cat = r["category"] or "未分类"
        if b not in tree:
            tree[b] = {}
            bank_counts[b] = 0
        if cat not in tree[b]:
            tree[b][cat] = []
        bank_counts[b] += 1
        tree[b][cat].append({
            "id": r["doc_id"],
            "title": r["title"] or "未知文档",
            "filename": r["filename"] or "",
            "doc_type": r["doc_type"] or "generic",
            "created": r["created_at"] or "",
        })
    
    # Get bank names from config
    bank_names = {k: v["name"] for k, v in BANKS.items() if k != "all"}
    
    return {
        "tree": tree,
        "bank_names": bank_names,
        "bank_counts": bank_counts,
        "total": len(rows),
    }

@app.get("/api/categories")
async def list_categories():
    """列出所有分类及文档数"""
    db = get_db()
    rows = db.execute(
        "SELECT category, COUNT(*) as cnt FROM doc_meta WHERE category != '' GROUP BY category ORDER BY cnt DESC"
    ).fetchall()
    db.close()

    used = {r["category"]: r["cnt"] for r in rows}
    result = []
    for cat in DEFAULT_CATEGORIES:
        result.append({"name": cat, "count": used.get(cat, 0)})
    for cat, cnt in used.items():
        if cat not in DEFAULT_CATEGORIES:
            result.append({"name": cat, "count": cnt})
    return {"categories": result}


@app.get("/api/banks")
async def list_banks():
    """列出所有 bank 及文档统计"""
    db = get_db()
    bank_stats = {}
    try:
        rows = db.execute("SELECT bank, COUNT(*) as cnt FROM doc_meta WHERE bank != 'skip' GROUP BY bank").fetchall()
        bank_stats = {r["bank"]: r["cnt"] for r in rows}
    except sqlite3.OperationalError:
        pass
    db.close()

    total = sum(bank_stats.get(key, 0) for key in BANKS if key != "all")
    banks = []
    for key, cfg in BANKS.items():
        if key == "all":
            banks.append({"key": key, "name": cfg["name"], "count": total})
        else:
            banks.append({"key": key, "name": cfg["name"], "count": bank_stats.get(key, 0)})
    return {"banks": banks}


@app.get("/api/stats")
async def stats():
    """知识库统计"""
    result = await hindsight_request(f"/v1/default/banks/{KB_BANK}/stats")
    return {
        "total_nodes": result.get("total_nodes", 0),
        "total_documents": result.get("total_documents", 0),
        "total_links": result.get("total_links", 0),
    }


@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str):
    """删除文档及其所有向量（按 doc_id tag 查找 Hindsight 内部文档）"""
    # 1. 获取所有 Hindsight 文档，找出带 doc_id:{doc_id} tag 的
    try:
        all_docs = await hindsight_request(
            f"/v1/default/banks/{KB_BANK}/documents?limit=500",
            timeout=15,
        )
    except Exception as e:
        print(f"获取 Hindsight 文档列表失败: {e}", file=sys.stderr)
        all_docs = {"items": []}
    
    # 2. 按 tag 匹配，删除所有关联的 Hindsight 内部文档
    deleted_hs = 0
    for item in all_docs.get("items", []):
        tags = item.get("tags", [])
        for t in tags:
            if t == f"doc_id:{doc_id}":
                try:
                    await hindsight_request(
                        f"/v1/default/banks/{KB_BANK}/documents/{item['id']}",
                        "DELETE",
                        timeout=10,
                    )
                    deleted_hs += 1
                except Exception as e:
                    print(f"删除 Hindsight 文档 {item['id'][:16]} 失败: {e}", file=sys.stderr)
                break  # 一个 Hindsight 文档只删一次
    
    # 3. 删除 SQLite 元数据
    delete_meta(doc_id)
    return {"ok": True, "deleted_hindsight_docs": deleted_hs}


@app.post("/api/reparse/{doc_id}")
async def reparse_document(doc_id: str):
    """重新解析已入库文档：删除旧向量 → 重新OCR → 重新入库
    
    用于升级到 MinerU 后对旧 tesseract 文档重新处理。
    需要文档对应的原始文件在 uploads/ 目录下。
    """
    # 查找文档元数据
    meta = get_meta(doc_id)
    if not meta or not meta.get("filename"):
        raise HTTPException(404, f"文档 {doc_id} 不存在或无文件名记录")
    
    filename = meta["filename"]
    doc_title = meta.get("title", filename_to_title(filename))
    doc_category = meta.get("category", "")
    
    # 查找原始文件
    upload_dir = os.path.join(BASE_DIR, "uploads")
    file_path = os.path.join(upload_dir, filename)
    
    if not os.path.exists(file_path):
        # 尝试直接从 Hindsight tags 重建（PDF 文件可能未保留）
        raise HTTPException(
            404,
            f"原始文件 {filename} 不在 uploads/ 目录中。"
            f"请重新上传文件以触发 MinerU 解析。"
        )
    
    # 读取文件
    with open(file_path, "rb") as f:
        content = f.read()
    
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(400, f"文件过大（{len(content)//1024//1024}MB），上限 {MAX_FILE_SIZE//1024//1024}MB")
    
    # 重新解析
    try:
        text = await parse_document(filename, content)
        # 清除空字节，PostgreSQL UTF8 不接受 0x00
        text = text.replace("\x00", "")
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        traceback.print_exc(file=sys.stderr)
        raise HTTPException(500, f"重新解析异常: {e}")
    
    if not text or len(text.strip()) < 10:
        raise HTTPException(400, "重新解析后文档内容过短")
    
    # 删除旧向量
    try:
        await hindsight_request(
            f"/v1/default/banks/{KB_BANK}/documents/{doc_id}",
            "DELETE",
            timeout=30,
        )
        print(f"已删除旧文档向量: {doc_id}", file=sys.stderr)
    except Exception as e:
        print(f"删除旧向量失败（继续）: {e}", file=sys.stderr)
    
    # 重新切片入库（复用 upload 后的逻辑）
    chunk_size = 1000
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    new_doc_id = str(uuid.uuid4())
    
    memory_items = []
    for i, chunk in enumerate(chunks):
        chunk = chunk.strip()
        if not chunk:
            continue
        tags = [
            f"doc:{filename}",
            f"chunk:{i+1}/{len(chunks)}",
            f"doc_id:{new_doc_id}",
            f"title:{doc_title}",
        ]
        if doc_category:
            tags.append(f"cat:{doc_category}")
        memory_items.append({"content": chunk, "tags": tags, "type": "world"})
    
    retained = 0
    if memory_items:
        dyn_timeout = max(120, min(len(memory_items) * 5, 600))
        print(f"Reparse: {len(memory_items)} chunks, timeout={dyn_timeout}s", file=sys.stderr)
        try:
            result = await hindsight_request(
                f"/v1/default/banks/{KB_BANK}/memories",
                "POST",
                {"items": memory_items},
                timeout=dyn_timeout,
            )
            retained = result.get("items_count", len(memory_items))
        except Exception as e:
            raise HTTPException(502, f"重新入库失败: {e}")
    
    # 更新元数据（保留旧标题和分类）
    content_hash = hashlib.sha256(content).hexdigest()
    save_meta(new_doc_id, doc_title, doc_category, filename, content_hash)
    
    return {
        "ok": True,
        "old_doc_id": doc_id,
        "new_doc_id": new_doc_id,
        "title": doc_title,
        "chunks": retained,
        "total_chars": len(text),
        "preview": text[:200] + ("..." if len(text) > 200 else ""),
    }


@app.get("/api/rag-eval")
async def rag_evaluation():
    """RAG 效果评估 — 4 维度（RAGAS 风格）"""

    # 预置测试用例：覆盖各 bank 典型查询
    test_cases = [
        {"q": "等保三级的安全区域边界要求是什么", "bank": "standards", "expect_doc": "等级保护"},
        {"q": "政务信息化项目验收流程有哪些步骤", "bank": "project_docs", "expect_doc": "验收"},
        {"q": "软件造价评估的方法和依据是什么", "bank": "industry_docs", "expect_doc": "造价"},
        {"q": "密码应用方案的测评要求是什么", "bank": "standards", "expect_doc": "密码"},
        {"q": "信息化项目立项咨询的主要内容", "bank": "project_docs", "expect_doc": "立项"},
        {"q": "等保测评的国家标准编号是什么", "bank": "standards", "expect_doc": "GB"},
        {"q": "政务信息化监理服务的职责范围", "bank": "standards", "expect_doc": "监理"},
        {"q": "数据中心建设的技术要求", "bank": "standards", "expect_doc": "数据中心"},
    ]

    results = []
    dimensions = {"retrieval": [], "groundedness": [], "relevance": [], "utilization": []}

    for tc in test_cases:
        try:
            # Step 1: 召回
            recalled = await recall(tc["q"], limit=10, bank="kb")
            context_texts = [r.get("text", "") for r in recalled if r.get("text")]
            context_block = "\n\n---\n\n".join(context_texts[:5])

            if not context_block.strip():
                results.append({"q": tc["q"], "bank": tc["bank"], "error": "无召回结果",
                               "scores": {"retrieval": 0, "groundedness": 0, "relevance": 0, "utilization": 0}})
                continue

            # Step 2: 生成答案（复用现有逻辑）
            bank_cfg = get_bank_config(tc["bank"])
            bank_prompt = bank_cfg["prompt"]
            answer_prompt = f"""{bank_prompt}

【硬性规则】
1. 只使用下方「文档内容」中的信息回答，禁止使用你的训练知识补充任何事实
2. 每个关键论断必须在括号内标注来源文档名称
3. 文档中没有相关信息时，直接回答「根据现有资料无法确定」

基于以下文档内容回答问题：

文档内容：
{context_block}

问题：{tc["q"]}

请用中文回答，引用具体条款和数据，并标注信息来源。"""

            answer = await deepseek_chat([
                {"role": "system", "content": bank_prompt},
                {"role": "user", "content": answer_prompt},
            ])

            # Step 3: LLM 评估 4 个维度
            eval_prompt = f"""你是一个 RAG 系统评估专家。请对以下问答对进行 4 个维度的评分。

【问题】{tc["q"]}

【检索到的文档片段】
{context_block[:3000]}

【生成的答案】
{answer[:2000]}

请对以下 4 个维度分别打分（0.0 ~ 1.0），并给出简短理由：

1. Retrieval（检索质量）：检索到的文档片段与问题的相关性
2. Groundedness（可对应性）：答案中的每句话能否在文档片段中找到依据
3. Relevance（答案相关性）：答案是否切题、是否回答了问题
4. Utilization（利用率）：答案是否充分利用了检索到的文档内容

请严格用以下 JSON 格式回复，不要有其他内容：
{{"retrieval": {{"score": 0.0, "reason": "..."}}, "groundedness": {{"score": 0.0, "reason": "..."}}, "relevance": {{"score": 0.0, "reason": "..."}}, "utilization": {{"score": 0.0, "reason": "..."}}}}"""

            eval_result = await deepseek_chat([
                {"role": "system", "content": "你是严格的 RAG 评估专家，只输出 JSON。"},
                {"role": "user", "content": eval_prompt},
            ])

            # 解析评估结果
            json_match = re.search(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', eval_result, re.DOTALL)
            scores = {"retrieval": 0, "groundedness": 0, "relevance": 0, "utilization": 0}
            if json_match:
                try:
                    parsed = json.loads(json_match.group())
                    for dim in ["retrieval", "groundedness", "relevance", "utilization"]:
                        val = parsed.get(dim, {})
                        if isinstance(val, dict):
                            scores[dim] = round(float(val.get("score", 0)), 2)
                        elif isinstance(val, (int, float)):
                            scores[dim] = round(float(val), 2)
                except Exception:
                    pass

            for dim in dimensions:
                dimensions[dim].append(scores[dim])

            results.append({
                "q": tc["q"],
                "bank": tc["bank"],
                "answer_preview": answer[:200],
                "chunks_recalled": len(recalled),
                "scores": scores,
                "eval_raw": eval_result[:500],
            })

        except Exception as e:
            results.append({"q": tc["q"], "bank": tc["bank"], "error": str(e)[:200],
                           "scores": {"retrieval": 0, "groundedness": 0, "relevance": 0, "utilization": 0}})

    # 汇总
    avg_scores = {}
    for dim in dimensions:
        vals = dimensions[dim]
        avg_scores[dim] = round(sum(vals) / max(len(vals), 1), 2)

    overall = round(sum(avg_scores.values()) / 4, 2)

    return {
        "total_cases": len(test_cases),
        "evaluated": len([r for r in results if "error" not in r]),
        "avg_scores": avg_scores,
        "overall": overall,
        "details": results,
    }


@app.get("/api/audit")
async def audit_knowledge_base():
    """扫描知识库所有文档，输出质量检查报告"""
    db = get_db()
    rows = db.execute(
        "SELECT doc_id, title, filename, bank, created_at FROM doc_meta WHERE bank NOT IN ('skip') ORDER BY created_at DESC"
    ).fetchall()
    db.close()

    results = []
    for row in rows:
        doc_id = row["doc_id"]
        title = row["title"] or "未知文档"
        bank = row["bank"] or "kb"

        # Get full text from Hindsight
        bank_cfg = BANKS.get(bank, BANKS["all"])
        hs_bank = bank_cfg["hindsight"]

        full_text = ""
        try:
            # 先按 doc_id 精确查命名 bank
            hindsight_doc_id = None
            docs_result = await hindsight_request(
                f"/v1/default/banks/{hs_bank}/documents", timeout=10
            )
            doc_list = docs_result.get("items", []) or docs_result.get("documents", [])
            for d in doc_list:
                if f"doc_id:{doc_id}" in d.get("tags", []):
                    hindsight_doc_id = d.get("id")
                    break
            
            # kb fallback
            if not hindsight_doc_id and hs_bank != "kb":
                docs_result = await hindsight_request(
                    f"/v1/default/banks/kb/documents", timeout=10
                )
                doc_list = docs_result.get("items", []) or docs_result.get("documents", [])
                for d in doc_list:
                    if f"doc_id:{doc_id}" in d.get("tags", []):
                        hindsight_doc_id = d.get("id")
                        hs_bank = "kb"
                        break
            
            if hindsight_doc_id:
                # 找到了：直接取 original_text
                doc_detail = await hindsight_request(
                    f"/v1/default/banks/{hs_bank}/documents/{hindsight_doc_id}", timeout=10
                )
                full_text = doc_detail.get("original_text", "") or doc_detail.get("text", "") or ""
            else:
                # doc_id 未找到（consolidation 导致），改用标题语义搜索
                try:
                    recall_result = await recall(title, limit=50, bank="kb", max_tokens=32768)
                    texts = []
                    for r in recall_result:
                        t = r.get("text", "") or ""
                        if t.strip():
                            texts.append(t.strip())
                    full_text = "\n\n".join(texts)
                except Exception:
                    pass
        except Exception:
            pass

        quality = assess_quality(full_text)

        # ── 完整性检查：对比原始文件 ──
        filename = row["filename"] or ""
        completeness = None  # {available, original_chars, retrieved_chars, coverage_pct}
        if filename:
            upload_path = os.path.join(BASE_DIR, "uploads", filename)
            if os.path.exists(upload_path):
                try:
                    with open(upload_path, "rb") as f:
                        raw = f.read()
                    # 尝试解析原文件获取字符数
                    orig_chars = 0
                    ext = os.path.splitext(filename)[1].lower()
                    if ext == ".pdf":
                        reader = pypdf.PdfReader(BytesIO(raw))
                        for page in reader.pages:
                            t = page.extract_text()
                            if t:
                                orig_chars += len(t)
                    elif ext in (".txt", ".md"):
                        orig_chars = len(raw.decode("utf-8", errors="ignore"))
                    elif ext == ".docx":
                        d = docx.Document(BytesIO(raw))
                        orig_chars = sum(len(p.text) for p in d.paragraphs)
                    
                    if orig_chars > 0:
                        coverage = min(100, round(len(full_text) / orig_chars * 100))
                        completeness = {
                            "available": True,
                            "original_chars": orig_chars,
                            "retrieved_chars": len(full_text),
                            "coverage_pct": coverage,
                        }
                except Exception:
                    pass
        
        if not completeness:
            completeness = {"available": False, "reason": "原文件不可用（已删除或未保留）"}

        results.append({
            "doc_id": doc_id,
            "title": title,
            "bank": bank,
            "chars": len(full_text),
            "score": quality["score"],
            "issues": quality["issues"],
            "needs_refetch": quality["score"] < 70,
            "completeness": completeness,
        })

    # Summary stats
    total = len(results)
    low_quality = [r for r in results if r["needs_refetch"]]
    avg_score = sum(r["score"] for r in results) / max(total, 1)

    return {
        "total_docs": total,
        "avg_score": round(avg_score, 1),
        "low_quality_count": len(low_quality),
        "documents": sorted(results, key=lambda x: x["score"]),
    }


@app.post("/api/audit/refetch")
async def refetch_document(doc_id: str = Form(...), std_no: str = Form("")):
    """重新下载标准文档，使用 MinerU 解析，替换旧数据"""
    import subprocess, tempfile, os as _os_module

    # Verify doc exists
    db = get_db()
    meta = db.execute("SELECT title, bank FROM doc_meta WHERE doc_id = ?", (doc_id,)).fetchone()
    if not meta:
        db.close()
        raise HTTPException(404, "文档不存在")

    old_bank = meta["bank"] or "kb"
    bank_cfg = BANKS.get(old_bank, BANKS["all"])
    hs_bank = bank_cfg["hindsight"]

    search_term = std_no.strip() or meta["title"]
    db.close()

    # Step 1: Search AnySearch
    anysearch_cli = _os_module.path.expanduser("~/.agents/skills/anysearch/scripts/anysearch_cli.py")
    try:
        result = subprocess.run(
            ["python3", anysearch_cli, "search", f"{search_term} 标准 PDF下载", "-m", "5", "--freshness", "year"],
            capture_output=True, text=True, timeout=30
        )
        urls = [line.split("**URL**: ")[1].strip() for line in result.stdout.split('\n') if "**URL**: " in line]
    except Exception as e:
        raise HTTPException(500, f"搜索失败: {e}")

    if not urls:
        raise HTTPException(404, f"未找到 {search_term} 的下载链接")

    # Step 2: Download PDF
    pdf_path = None
    pdf_content = None
    async with httpx.AsyncClient(timeout=90, follow_redirects=True) as client:
        for url in urls[:3]:
            try:
                resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
                if resp.status_code == 200:
                    ct = resp.headers.get("content-type", "")
                    if "application/pdf" in ct or url.lower().endswith(".pdf"):
                        pdf_content = resp.content
                        fd, pdf_path = tempfile.mkstemp(suffix=".pdf")
                        _os_module.write(fd, pdf_content)
                        _os_module.close(fd)
                        break
            except Exception:
                continue

    if not pdf_path or not pdf_content:
        raise HTTPException(500, "下载失败")

    # Step 3: Parse with MinerU
    text = ""
    try:
        text = await mineru_parse_pdf(f"{search_term}.pdf", pdf_content)
    except Exception:
        # Fallback to pypdf
        with open(pdf_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"

    _os_module.unlink(pdf_path)

    if not text or len(text.strip()) < 100:
        raise HTTPException(400, "PDF解析后内容过短")

    # Step 4: Delete old vectors from Hindsight
    try:
        docs_result = await hindsight_request(
            f"/v1/default/banks/{hs_bank}/documents", timeout=10
        )
        doc_list = docs_result.get("items", []) or docs_result.get("documents", [])
        for d in doc_list:
            tags = d.get("tags", [])
            if f"doc_id:{doc_id}" in tags:
                try:
                    await hindsight_request(
                        f"/v1/default/banks/{hs_bank}/documents/{d['id']}",
                        method="DELETE", timeout=10
                    )
                except Exception:
                    pass
    except Exception:
        pass

    # Step 5: Re-upload with new text
    chunk_size = 1000
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

    success_count = 0
    total_count = 0
    async with httpx.AsyncClient(timeout=30) as client:
        for i, chunk in enumerate(chunks):
            chunk = chunk.strip()
            if not chunk:
                continue
            total_count += 1
            payload = [{
                "text": chunk,
                "tags": [
                    f"doc:{search_term}.pdf",
                    f"chunk:{i+1}/{len(chunks)}",
                    f"doc_id:{doc_id}",
                    f"title:{search_term}",
                ]
            }]
            try:
                r = await client.post(
                    f"{HINDSIGHT_URL}/v1/default/banks/{hs_bank}/memories",
                    json=payload
                )
                if r.status_code in (200, 201):
                    success_count += 1
            except Exception:
                pass

    # Update meta.db title
    db = get_db()
    db.execute("UPDATE doc_meta SET title = ? WHERE doc_id = ?", (search_term, doc_id))
    db.commit()
    db.close()

    quality = assess_quality(text)

    return {
        "ok": True,
        "doc_id": doc_id,
        "title": search_term,
        "text_length": len(text),
        "chunks": f"{success_count}/{total_count}",
        "new_score": quality["score"],
        "used_mineru": bool(MINERU_TOKEN),
    }


@app.get("/", response_class=HTMLResponse)
async def index():
    return HTMLResponse(INDEX_HTML)


# ─── 前端 ─────────────────────────────────────────────────────

INDEX_HTML = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>知识库</title>
<style>
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;background:#f5f5f5;color:#222;min-height:100vh}
.header{background:#1a1a2e;color:#fff;padding:14px 24px;display:flex;align-items:center;justify-content:space-between}
.header h1{font-size:17px;font-weight:600}
.stats{font-size:12px;color:#aab;display:flex;gap:14px}
.bank-bar{display:flex;background:#fff;border-bottom:1px solid #e0e0e0;padding:8px 16px;align-items:center;gap:8px}
.bank-bar label{font-size:12px;color:#888;white-space:nowrap}
.bank-bar select{padding:6px 10px;border:1px solid #ddd;border-radius:6px;font-size:13px;outline:none;background:#fff;min-width:120px}
.bank-bar select:focus{border-color:#e94560}
.doc-bank-select{font-size:11px;padding:2px 4px;border:1px solid #ddd;border-radius:4px;background:#fff;color:#666;outline:none;max-width:80px}
.doc-bank-select:focus{border-color:#e94560}
.main-area{max-width:900px;margin:0 auto;padding:20px}
.search-box{display:flex;gap:8px;margin-bottom:16px}
.search-box input{flex:1;padding:12px 16px;border:1px solid #ddd;border-radius:8px;font-size:15px;outline:none}
.search-box input:focus{border-color:#e94560}
.search-box button{padding:12px 24px;background:#e94560;color:#fff;border:none;border-radius:8px;font-size:15px;cursor:pointer;font-weight:600}
.search-box button:hover{background:#d63850}
.answer{background:#fff;border-radius:10px;padding:20px;line-height:1.8;font-size:15px;margin-bottom:14px;box-shadow:0 1px 3px rgba(0,0,0,.05)}
.answer .bank-label{display:inline-block;padding:2px 10px;border-radius:10px;font-size:11px;margin-bottom:10px;color:#fff}
.answer .bank-label.all{background:#666}
.answer .bank-label.tech{background:#8e44ad}
.answer .bank-label.security{background:#c0392b}
.answer .bank-label.ai{background:#2980b9}
.answer .bank-label.notes{background:#16a085}
.answer .bank-label.proposals{background:#4a90d9}
.answer .bank-label.assessment{background:#e67e22}
.answer .bank-label.projects{background:#27ae60}
.sources{margin-top:16px;padding-top:12px;border-top:1px solid #eee}
.source{background:#f8f8f8;border-radius:6px;padding:10px 14px;margin-bottom:8px;font-size:13px}
.source .doc{color:#e94560;font-weight:600;margin-bottom:4px}
.source .text{color:#666;line-height:1.6}
.section-bar{display:flex;gap:8px;margin-bottom:16px;align-items:center}
.section-btn{padding:8px 16px;border-radius:6px;font-size:13px;cursor:pointer;border:1px solid #ddd;background:#fff;color:#666;transition:.2s}
.section-btn:hover,.section-btn.active{border-color:#e94560;color:#e94560}
.upload-form{background:#fff;border-radius:12px;padding:24px;margin-bottom:16px;display:none;box-shadow:0 1px 3px rgba(0,0,0,.05)}
.upload-form.show{display:block}
.standard-form{background:#fff;border-radius:12px;padding:24px;margin-bottom:16px;display:none;box-shadow:0 1px 3px rgba(0,0,0,.05)}
.standard-form.show{display:block}
.audit-form{background:#fff;border-radius:12px;padding:24px;margin-bottom:16px;display:none;box-shadow:0 1px 3px rgba(0,0,0,.05)}
.audit-form.show{display:block}
.rag-eval-form{background:#fff;border-radius:12px;padding:24px;margin-bottom:16px;display:none;box-shadow:0 1px 3px rgba(0,0,0,.05)}
.rag-eval-form.show{display:block}
.form-group{margin-bottom:14px}
.form-group label{display:block;font-size:13px;font-weight:600;color:#555;margin-bottom:6px}
.form-group input,.form-group select{width:100%;padding:12px 14px;border:1px solid #ddd;border-radius:8px;font-size:14px;outline:none}
.form-group select{appearance:none;background:#fff url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='12' height='12' viewBox='0 0 12 12'%3E%3Cpath fill='%23888' d='M6 8L1 3h10z'/%3E%3C/svg%3E") no-repeat right 14px center;padding-right:36px}
.upload-zone{border:2px dashed #ccc;border-radius:12px;padding:30px;text-align:center;cursor:pointer;transition:.2s;margin-bottom:16px}
.upload-zone:hover,.upload-zone.dragover{border-color:#e94560;background:#fff5f5}
.upload-zone .icon{font-size:36px;margin-bottom:8px}
.upload-zone .hint{color:#999;font-size:13px}
.upload-zone .file-selected{color:#e94560;font-weight:600;font-size:14px}
.upload-zone input[type=file]{display:none}
.upload-btn{width:100%;padding:14px;background:#e94560;color:#fff;border:none;border-radius:8px;font-size:15px;cursor:pointer;font-weight:600}
.upload-btn:hover{background:#d63850}
.upload-btn:disabled{background:#ccc;cursor:not-allowed}
.result-msg{margin-top:12px;padding:12px;border-radius:6px;font-size:14px}
.result-msg.success{background:#e8f5e9;color:#2e7d32}
.result-msg.error{background:#ffebee;color:#c62828}
.doc-list{background:#fff;border-radius:10px;display:none;box-shadow:0 1px 3px rgba(0,0,0,.05)}
.doc-list.show{display:block}
.doc-item{display:flex;justify-content:space-between;align-items:flex-start;padding:14px 16px;border-bottom:1px solid #f0f0f0;gap:12px}
.doc-item:last-child{border-bottom:none}
.doc-item .main{flex:1;min-width:0}
.doc-item .title-line{display:flex;align-items:center;gap:8px;margin-bottom:4px}
.doc-item .title{font-weight:600;font-size:14px;color:#222}
.doc-item .cat-tag{padding:2px 8px;border-radius:10px;font-size:11px;background:#f0f0f0;color:#888;white-space:nowrap}
.doc-item .meta{color:#999;font-size:12px}
.doc-item .actions{display:flex;gap:6px;flex-shrink:0}
.doc-item .actions span{cursor:pointer;padding:4px 8px;border-radius:4px;font-size:13px}
.doc-item .del{color:#c62828}
.doc-item .del:hover{background:#ffebee}
.loading{text-align:center;padding:20px;color:#999}
.empty{text-align:center;padding:40px;color:#bbb;font-size:14px}
.spin{display:inline-block;width:16px;height:16px;border:2px solid #ddd;border-top-color:#e94560;border-radius:50%;animation:spin .6s linear infinite;margin-right:8px}
@keyframes spin{to{transform:rotate(360deg)}}
.wiki-bank{margin-bottom:12px;border:1px solid #eee;border-radius:8px;overflow:hidden}
.wiki-bank-header{display:flex;align-items:center;padding:12px 16px;cursor:pointer;user-select:none;background:#fafafa;border-bottom:1px solid #eee;transition:.2s}
.wiki-bank-header:hover{background:#f0f0f0}
.wiki-bank-header .icon{margin-right:8px;font-size:16px}
.wiki-bank-header .name{font-weight:600;font-size:14px;flex:1}
.wiki-bank-header .count{color:#888;font-size:12px;padding:2px 8px;background:#f0f0f0;border-radius:10px}
.wiki-bank-header .arrow{color:#aaa;transition:.2s;font-size:12px}
.wiki-bank.open>.wiki-bank-header .arrow{transform:rotate(90deg)}
.wiki-bank-body{display:none;padding:8px 0}
.wiki-bank.open>.wiki-bank-body{display:block}
.wiki-cat{padding:4px 0}
.wiki-cat-header{display:flex;align-items:center;padding:8px 16px;cursor:pointer;font-size:13px;color:#666;transition:.2s}
.wiki-cat-header:hover{background:#f8f8f8}
.wiki-cat-header .cat-icon{margin-right:6px}
.wiki-cat-header .cat-name{flex:1;font-weight:500}
.wiki-cat-header .cat-count{color:#aaa;font-size:11px}
.wiki-cat-header .cat-arrow{color:#ccc;font-size:11px;margin-left:4px;transition:.2s}
.wiki-cat.open>.wiki-cat-header .cat-arrow{transform:rotate(90deg)}
.wiki-cat-body{display:none;padding-left:32px}
.wiki-cat.open>.wiki-cat-body{display:block}
.wiki-doc{padding:8px 16px;font-size:13px;color:#444;display:flex;align-items:center;gap:8px;cursor:pointer;border-radius:4px;margin:2px 8px;transition:.15s}
.wiki-doc:hover{background:#f5f5f5;color:#e94560}
.wiki-doc .doc-icon{font-size:14px}
.wiki-doc .doc-title{flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.wiki-doc .doc-type{font-size:10px;padding:1px 6px;border-radius:8px;background:#f0f0f0;color:#888}
.wiki-doc-content{background:#f8f8f8;border-radius:8px;padding:16px;margin:8px 16px 16px 48px;font-size:13px;line-height:1.8;white-space:pre-wrap;max-height:500px;overflow-y:auto;display:none;border:1px solid #eee}
.wiki-doc-content.show{display:block}
</style>
</head>
<body>
<div class="header" style="display:flex;align-items:center;justify-content:space-between">
  <h1>📚 知识库</h1>
  <div style="display:flex;align-items:center;gap:14px">
    <button onclick="resetAll()" style="background:none;border:1px solid rgba(255,255,255,0.3);color:#fff;padding:5px 12px;border-radius:6px;cursor:pointer;font-size:13px" title="刷新页面">🔄 刷新</button>
    <div class="stats" id="stats">加载中...</div>
  </div>
</div>
<div class="bank-bar" id="bank-bar">
  <label>📂 知识库分类：</label>
  <select id="bank-selector" onchange="switchBank(this.value)"></select>
</div>

<div class="main-area">
  <div class="search-box">
    <input id="query" placeholder="输入问题搜索知识库..." onkeydown="if(event.key==='Enter')doSearch()">
    <button onclick="doSearch()">搜索</button>
  </div>
  <div id="answer-area"></div>

  <div class="section-bar">
    <button class="section-btn active" id="btn-upload" onclick="toggleSection('upload')">📤 上传文档</button>
    <button class="section-btn" id="btn-docs" onclick="toggleSection('docs')">📋 文档列表</button>
    <button class="section-btn" id="btn-standard" onclick="toggleSection('standard')">📥 规范下载</button>
    <button class="section-btn" id="btn-audit" onclick="toggleSection('audit')">🔍 数据自查</button>
    <button class="section-btn" id="btn-rag-eval" onclick="toggleSection('rag-eval')">📊 RAG评估</button>
    <button class="section-btn" id="btn-wiki" onclick="toggleSection('wiki')">📚 知识浏览</button>
  </div>

  <div class="upload-form" id="upload-form">
    <div class="form-group">
      <label>📌 文档标题 <span style="color:#999;font-weight:400">（可选，默认使用文件名）</span></label>
      <input id="upload-title" placeholder="输入文档标题...">
    </div>
    <div class="form-group">
      <label>📂 分类</label>
      <select id="upload-category"><option value="">-- 未分类 --</option></select>
    </div>
    <div class="form-group">
      <label>🏷️ 知识库分类</label>
      <select id="upload-bank"><option value="">-- 加载中 --</option></select>
    </div>
    <div class="upload-zone" id="drop-zone" onclick="document.getElementById('file-input').click()">
      <div class="icon">📄</div>
      <div class="hint" id="drop-hint">点击选择文件 / 拖拽上传<br>支持 PDF / Word / Markdown / TXT<br><small style="color:#aaa">支持多选和文件夹</small></div>
      <input type="file" id="file-input" accept=".pdf,.docx,.doc,.md,.txt" multiple onchange="onFilesSelected(this.files)">
      <input type="file" id="folder-input" webkitdirectory onchange="onFilesSelected(this.files)" style="display:none">
    </div>
    <div style="display:flex;gap:8px;margin-bottom:16px">
      <button class="upload-btn" id="upload-btn" disabled onclick="doUpload()" style="flex:1">上传到知识库</button>
      <button class="upload-btn" id="folder-btn" onclick="document.getElementById('folder-input').click()" style="flex:0 0 auto;background:#555;font-size:13px;padding:14px 18px">📁 文件夹</button>
    </div>
    <div id="upload-result"></div>
  </div>

  <div class="standard-form" id="standard-form">
    <div class="form-group">
      <label>📌 标准号</label>
      <input id="std-no" placeholder="例如: GB 50348, GB/T 22239, ISO 27001" style="width:100%;padding:12px 14px;border:1px solid #ddd;border-radius:8px;font-size:14px;outline:none" onkeydown="if(event.key==='Enter')doFetchStandard()">
    </div>
    <div class="form-group">
      <label>🏷️ 入库分类</label>
      <select id="std-bank"><option value="">-- 加载中 --</option></select>
    </div>
    <button class="upload-btn" onclick="doFetchStandard()">🔍 搜索并下载入库</button>
    <div id="std-result" style="margin-top:12px"></div>
  </div>

  <div class="audit-form" id="audit-form">
    <button class="upload-btn" onclick="doAudit()" style="margin-bottom:14px">🔍 开始自查</button>
    <div id="audit-result"></div>
  </div>

  <div class="rag-eval-form" id="rag-eval-form">
    <div style="margin-bottom:12px;">
      <h3 style="margin:0 0 6px;">📊 RAG 效果评估（4维度）</h3>
      <p style="color:#888;font-size:13px;margin:0;">基于 RAGAS 框架评估检索质量、可对应性、答案相关性和利用率</p>
    </div>
    <button class="upload-btn" onclick="loadRagEval()" style="margin-bottom:14px;background:#C85032;">🚀 开始评估</button>
    <div id="rag-eval-status" style="margin-top:12px;color:#888;font-size:13px;"></div>
    <div id="rag-eval-results" style="margin-top:16px;"></div>
  </div>

  <div id="wiki-section" style="display:none">
    <div style="background:#fff;border-radius:12px;padding:20px;box-shadow:0 1px 3px rgba(0,0,0,.05)">
      <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:16px">
        <h3 style="margin:0;font-size:16px">📚 知识库目录</h3>
        <span id="wiki-count" style="color:#888;font-size:12px"></span>
      </div>
      <div id="wiki-tree" class="wiki-tree"></div>
    </div>
  </div>

  <div id="doc-search-bar" style="display:none;margin-bottom:12px">
    <input id="doc-search" placeholder="🔍 搜索文档标题..." oninput="renderDocs()" style="width:100%;padding:10px 14px;border:1px solid #ddd;border-radius:8px;font-size:14px;outline:none">
  </div>
  <div class="doc-list" id="doc-list"></div>
</div>

<script>
let currentBank = 'all';
let selectedFile = null;
let allDocs = [];
let bankData = [];

async function init() {
  await loadBanks();
  await loadCategories();
  loadStats();
}
init();

async function loadBanks() {
  try {
    const r = await fetch('/api/banks');
    bankData = (await r.json()).banks;
    // Populate main bank selector
    const sel = document.getElementById('bank-selector');
    sel.innerHTML = bankData.map(b => 
      `<option value="${b.key}"${b.key === 'all' ? ' selected' : ''}>${b.name} (${b.count || 0})</option>`
    ).join('');
    // Populate upload bank selector
    const upSel = document.getElementById('upload-bank');
    upSel.innerHTML = '<option value="">-- 选择分类 --</option>' +
      bankData.filter(b => b.key !== 'all').map(b =>
        `<option value="${b.key}">${b.name} (${b.count || 0})</option>`
      ).join('');
    // Populate doc row bank dropdowns if visible
    if (document.getElementById('doc-list').classList.contains('show')) {
      populateDocBankSelects();
    }
  } catch(e) {}
}

function switchBank(bank) {
  currentBank = bank;
  document.getElementById('bank-selector').value = bank;
  document.getElementById('answer-area').innerHTML = '';
  document.getElementById('query').value = '';
  document.getElementById('query').focus();
  if (document.getElementById('doc-list').classList.contains('show')) {
    loadDocs();
  }
}

function resetAll() {
  currentBank = 'all';
  document.getElementById('bank-selector').value = 'all';
  document.getElementById('answer-area').innerHTML = '';
  document.getElementById('query').value = '';
  document.getElementById('query').focus();
}

async function doSearch() {
  const q = document.getElementById('query').value.trim();
  if (!q) return;
  const area = document.getElementById('answer-area');
  area.innerHTML = '<div class="loading"><span class="spin"></span>搜索中...</div>';
  
  try {
    const fd = new FormData(); fd.append('q', q); fd.append('bank', currentBank);
    const r = await fetch('/api/query', {method:'POST', body:fd});
    const d = await r.json();
    const bankName = (bankData.find(b => b.key === currentBank) || {}).name || currentBank;
    
    let html = '<div class="answer"><span class="bank-label '+currentBank+'">'+bankName+'</span><br>' + d.answer.replace(/\\\\\\n/g,'<br>') + '</div>';
    if (d.sources && d.sources.length) {
      html += '<div class="sources"><strong>📎 参考来源</strong></div>';
      d.sources.forEach((s, idx) => {
        const docId = s.doc_id || '';
        const viewLink = docId ? ` <a href="javascript:viewDocContent('${docId}','${s.doc.replace(/'/g,"\\\\'")}'" style="color:#e94560;font-size:11px">📖 查看原文</a>` : '';
        html += `<div class="source" id="src-${idx}">
          <div class="doc">${s.doc} · ${s.chunk}${viewLink}</div>
          <div class="text">${s.text}</div>
          <div id="src-content-${idx}" style="display:none;margin-top:8px;padding:8px;background:#fff;border:1px solid #eee;border-radius:4px;max-height:400px;overflow-y:auto;font-size:12px;line-height:1.8;white-space:pre-wrap"></div>
        </div>`;
      });
    }
    // 联网搜索按钮
    html += '<div style="margin-top:12px;text-align:center"><button onclick="doWebSearch()" style="padding:8px 20px;background:#f0f0f0;color:#555;border:1px solid #ddd;border-radius:8px;font-size:13px;cursor:pointer">🌐 联网搜索</button></div>';
    area.innerHTML = html;
  } catch(e) {
    area.innerHTML = '<div class="result-msg error">搜索失败: '+e.message+'</div>';
  }
}

async function viewDocContent(docId, title) {
  // Find the source container
  const allSources = document.querySelectorAll('.source');
  let targetIdx = -1;
  allSources.forEach((src, i) => {
    if (src.querySelector('a') && src.querySelector('a').href && src.querySelector('a').href.includes(docId)) {
      targetIdx = i;
    }
  });
  if (targetIdx < 0) {
    // Fallback: find by docId in the href
    for (let i = 0; i < allSources.length; i++) {
      const a = allSources[i].querySelector('a');
      if (a && a.getAttribute('href') && a.getAttribute('href').includes(docId)) {
        targetIdx = i;
        break;
      }
    }
  }
  if (targetIdx < 0) return;
  
  const contentDiv = document.getElementById('src-content-' + targetIdx);
  if (contentDiv.style.display === 'block') {
    contentDiv.style.display = 'none';
    return;
  }
  
  contentDiv.style.display = 'block';
  contentDiv.innerHTML = '<div class="loading"><span class="spin"></span>加载原文...</div>';
  
  try {
    const r = await fetch('/api/documents/' + docId + '/content');
    const d = await r.json();
    if (!d.text) {
      contentDiv.innerHTML = '<div class="result-msg error">' + (d.detail || '文档内容未找到') + '</div>';
      return;
    }
    const NL = String.fromCharCode(10);
    const safeTitle = (d.title || '未知').replace(/</g,'&lt;');
    const safeText = d.text.replace(/</g,'&lt;').replace(/\\\\\\\\n/g, NL);
    contentDiv.innerHTML = '<strong>📄 ' + safeTitle + '</strong>（' + d.chunks + ' 个片段，' + d.text.length + ' 字符）<hr style="border-color:#eee;margin:6px 0">' + safeText.replace(new RegExp(NL,'g'),'<br>');
  } catch(e) {
    contentDiv.innerHTML = '<div class="result-msg error">加载失败: '+e.message+'</div>';
  }
}

function toggleSection(name) {
  if (name === 'upload') {
    const form = document.getElementById('upload-form');
    form.classList.toggle('show');
    document.getElementById('btn-upload').classList.toggle('active', form.classList.contains('show'));
    if (form.classList.contains('show')) {
      document.getElementById('doc-list').classList.remove('show');
      document.getElementById('standard-form').classList.remove('show');
      document.getElementById('audit-form').classList.remove('show');
      document.getElementById('rag-eval-form').classList.remove('show');
      document.getElementById('btn-docs').classList.remove('active');
      document.getElementById('btn-standard').classList.remove('active');
      document.getElementById('btn-audit').classList.remove('active');
      document.getElementById('btn-rag-eval').classList.remove('active');
      document.getElementById('wiki-section').style.display = 'none';
      document.getElementById('btn-wiki').classList.remove('active');
    }
  } else if (name === 'docs') {
    const list = document.getElementById('doc-list');
    list.classList.toggle('show');
    document.getElementById('btn-docs').classList.toggle('active', list.classList.contains('show'));
    if (list.classList.contains('show')) {
      document.getElementById('upload-form').classList.remove('show');
      document.getElementById('standard-form').classList.remove('show');
      document.getElementById('audit-form').classList.remove('show');
      document.getElementById('rag-eval-form').classList.remove('show');
      document.getElementById('btn-upload').classList.remove('active');
      document.getElementById('btn-standard').classList.remove('active');
      document.getElementById('btn-audit').classList.remove('active');
      document.getElementById('btn-rag-eval').classList.remove('active');
      document.getElementById('wiki-section').style.display = 'none';
      document.getElementById('btn-wiki').classList.remove('active');
      document.getElementById('doc-search-bar').style.display = 'block';
      loadDocs();
    } else {
      document.getElementById('doc-search-bar').style.display = 'none';
    }
  } else if (name === 'standard') {
    const stdForm = document.getElementById('standard-form');
    stdForm.classList.toggle('show');
    document.getElementById('btn-standard').classList.toggle('active', stdForm.classList.contains('show'));
    if (stdForm.classList.contains('show')) {
      document.getElementById('upload-form').classList.remove('show');
      document.getElementById('doc-list').classList.remove('show');
      document.getElementById('audit-form').classList.remove('show');
      document.getElementById('rag-eval-form').classList.remove('show');
      document.getElementById('btn-upload').classList.remove('active');
      document.getElementById('btn-docs').classList.remove('active');
      document.getElementById('btn-audit').classList.remove('active');
      document.getElementById('btn-rag-eval').classList.remove('active');
      document.getElementById('wiki-section').style.display = 'none';
      document.getElementById('btn-wiki').classList.remove('active');
      document.getElementById('doc-search-bar').style.display = 'none';
      // Populate bank dropdown for standard form
      const stdBank = document.getElementById('std-bank');
      if (stdBank && bankData.length > 0 && stdBank.options.length <= 1) {
        stdBank.innerHTML = '<option value="">-- 选择分类 --</option>' +
          bankData.filter(b => b.key !== 'all').map(b => '<option value="' + b.key + '">' + b.name + '</option>').join('');
      }
    }
  } else if (name === 'audit') {
    const auditForm = document.getElementById('audit-form');
    auditForm.classList.toggle('show');
    document.getElementById('btn-audit').classList.toggle('active', auditForm.classList.contains('show'));
    if (auditForm.classList.contains('show')) {
      document.getElementById('upload-form').classList.remove('show');
      document.getElementById('doc-list').classList.remove('show');
      document.getElementById('standard-form').classList.remove('show');
      document.getElementById('rag-eval-form').classList.remove('show');
      document.getElementById('btn-upload').classList.remove('active');
      document.getElementById('btn-docs').classList.remove('active');
      document.getElementById('btn-standard').classList.remove('active');
      document.getElementById('btn-rag-eval').classList.remove('active');
      document.getElementById('wiki-section').style.display = 'none';
      document.getElementById('btn-wiki').classList.remove('active');
      document.getElementById('doc-search-bar').style.display = 'none';
      doAudit();
    }
  } else if (name === 'rag-eval') {
    const ragForm = document.getElementById('rag-eval-form');
    ragForm.classList.toggle('show');
    document.getElementById('btn-rag-eval').classList.toggle('active', ragForm.classList.contains('show'));
    if (ragForm.classList.contains('show')) {
      document.getElementById('upload-form').classList.remove('show');
      document.getElementById('doc-list').classList.remove('show');
      document.getElementById('standard-form').classList.remove('show');
      document.getElementById('audit-form').classList.remove('show');
      document.getElementById('btn-upload').classList.remove('active');
      document.getElementById('btn-docs').classList.remove('active');
      document.getElementById('btn-standard').classList.remove('active');
      document.getElementById('btn-audit').classList.remove('active');
      document.getElementById('wiki-section').style.display = 'none';
      document.getElementById('btn-wiki').classList.remove('active');
      document.getElementById('doc-search-bar').style.display = 'none';
    }
  } else if (name === 'wiki') {
    const wikiSection = document.getElementById('wiki-section');
    wikiSection.style.display = wikiSection.style.display === 'none' ? 'block' : 'none';
    document.getElementById('btn-wiki').classList.toggle('active', wikiSection.style.display === 'block');
    if (wikiSection.style.display === 'block') {
      document.getElementById('upload-form').classList.remove('show');
      document.getElementById('doc-list').classList.remove('show');
      document.getElementById('standard-form').classList.remove('show');
      document.getElementById('audit-form').classList.remove('show');
      document.getElementById('rag-eval-form').classList.remove('show');
      document.getElementById('btn-upload').classList.remove('active');
      document.getElementById('btn-docs').classList.remove('active');
      document.getElementById('btn-standard').classList.remove('active');
      document.getElementById('btn-audit').classList.remove('active');
      document.getElementById('btn-rag-eval').classList.remove('active');
      document.getElementById('doc-search-bar').style.display = 'none';
      loadWiki();
    }
  }
}

async function loadRagEval() {
  const status = document.getElementById('rag-eval-status');
  const results = document.getElementById('rag-eval-results');
  status.textContent = '⏳ 正在评估中（约1-2分钟）...';
  results.innerHTML = '';
  try {
    const r = await fetch('/api/rag-eval');
    const d = await r.json();
    if (d.error) { status.textContent = '❌ ' + d.error; return; }

    let html = '<div style="background:#f5f5f0;border-radius:8px;padding:16px;margin-bottom:12px;">';
    html += '<div style="font-size:20px;font-weight:bold;margin-bottom:8px;">综合评分：' + (d.overall * 100).toFixed(0) + '%</div>';
    html += '<div style="display:flex;gap:16px;flex-wrap:wrap;">';
    const dimNames = {retrieval:'检索质量',groundedness:'可对应性',relevance:'答案相关性',utilization:'利用率'};
    const dimColors = {retrieval:'#4CAF50',groundedness:'#2196F3',relevance:'#FF9800',utilization:'#9C27B0'};
    for (const [k,v] of Object.entries(d.avg_scores)) {
      const pct = (v * 100).toFixed(0);
      html += '<div style="text-align:center;min-width:100px;">';
      html += '<div style="font-size:24px;font-weight:bold;color:' + (dimColors[k]||'#333') + ';">' + pct + '%</div>';
      html += '<div style="font-size:12px;color:#666;">' + (dimNames[k]||k) + '</div>';
      html += '</div>';
    }
    html += '</div>';
    html += '<div style="font-size:12px;color:#888;margin-top:8px;">测试用例：' + d.total_cases + ' 个 | 成功评估：' + d.evaluated + ' 个</div>';
    html += '</div>';

    // 详细结果
    html += '<div style="font-size:14px;font-weight:bold;margin:12px 0 8px;">详细结果</div>';
    for (const item of d.details) {
      const hasErr = !!item.error;
      const bgColor = hasErr ? '#fff0f0' : '#fafafa';
      html += '<div style="background:' + bgColor + ';border-radius:6px;padding:10px;margin-bottom:8px;font-size:13px;">';
      html += '<div style="font-weight:bold;">' + (item.q||'') + ' <span style="color:#888;font-weight:normal;">[' + (item.bank||'') + ']</span></div>';
      if (hasErr) {
        html += '<div style="color:red;">错误：' + item.error + '</div>';
      } else {
        const s = item.scores || {};
        html += '<div style="margin-top:4px;">检索:' + ((s.retrieval||0)*100).toFixed(0) + '% · 可对应:' + ((s.groundedness||0)*100).toFixed(0) + '% · 相关:' + ((s.relevance||0)*100).toFixed(0) + '% · 利用:' + ((s.utilization||0)*100).toFixed(0) + '%</div>';
        html += '<div style="color:#666;margin-top:4px;">召回 ' + (item.chunks_recalled||0) + ' 条 | ' + (item.answer_preview||'').substring(0,100) + '...</div>';
      }
      html += '</div>';
    }

    results.innerHTML = html;
    status.textContent = '';
  } catch(e) {
    status.textContent = '❌ 请求失败：' + e.message;
  }
}
let wikiData = null;

async function loadWiki() {
  const treeDiv = document.getElementById('wiki-tree');
  const countSpan = document.getElementById('wiki-count');
  treeDiv.innerHTML = '<div class="loading"><span class="spin"></span>加载目录...</div>';
  
  try {
    const r = await fetch('/api/wiki?bank=' + encodeURIComponent(currentBank));
    const data = await r.json();
    wikiData = data;
    countSpan.textContent = '共 ' + data.total + ' 个文档';
    
    if (data.total === 0) {
      treeDiv.innerHTML = '<div class="empty">📭 暂无文档</div>';
      return;
    }
    
    let html = '';
    const bankOrder = ['kb', 'standards', 'industry_docs', 'tech', 'proposals', 'assessment', 'ai', 'notes', 'projects', 'general'];
    const bankIcons = {'kb':'📚','standards':'📏','industry_docs':'📄','tech':'💻','proposals':'📝','assessment':'🔍','ai':'🤖','notes':'📝','projects':'📋','general':'📦'};
    
    // Show banks in order, then any remaining
    const shownBanks = new Set();
    for (const b of bankOrder) {
      if (!data.tree[b]) continue;
      shownBanks.add(b);
      const bankName = data.bank_names[b] || b;
      const count = data.bank_counts[b] || 0;
      const icon = bankIcons[b] || '📁';
      
      html += '<div class="wiki-bank" id="wiki-bank-' + b + '">';
      html += '<div class="wiki-bank-header" data-bank="' + b + '">';
      html += '<span class="icon">' + icon + '</span>';
      html += '<span class="name">' + bankName + '</span>';
      html += '<span class="count">' + count + ' 文档</span>';
      html += '<span class="arrow">▶</span>';
      html += '</div>';
      html += '<div class="wiki-bank-body">';
      
      const cats = data.tree[b];
      const catNames = Object.keys(cats).sort();
      for (const cat of catNames) {
        const docs = cats[cat];
        html += '<div class="wiki-cat" id="wiki-cat-' + b + '-' + cat.replace(/[^a-zA-Z0-9\u4e00-\u9fff]/g, '') + '">';
        html += '<div class="wiki-cat-header" data-cat="' + cat + '">';
        html += '<span class="cat-icon">📁</span>';
        html += '<span class="cat-name">' + cat + '</span>';
        html += '<span class="cat-count">' + docs.length + '</span>';
        html += '<span class="cat-arrow">▶</span>';
        html += '</div>';
        html += '<div class="wiki-cat-body">';
        
        for (const doc of docs) {
          const typeLabel = doc.doc_type === 'gb_standard' ? 'GB标准' : 
                           doc.doc_type === 'regulation' ? '法规' : '';
          html += '<div class="wiki-doc" data-docid="' + doc.id + '">';
          html += '<span class="doc-icon">' + (doc.doc_type === 'gb_standard' ? '📏' : doc.doc_type === 'regulation' ? '⚖️' : '📄') + '</span>';
          html += '<span class="doc-title">' + doc.title.replace(/</g, '&lt;') + '</span>';
          if (typeLabel) html += '<span class="doc-type">' + typeLabel + '</span>';
          html += '</div>';
          html += '<div class="wiki-doc-content" id="wiki-doc-' + doc.id + '"></div>';
        }
        
        html += '</div></div>';
      }
      
      html += '</div></div>';
    }
    
    // Show any remaining banks not in the predefined order
    for (const b of Object.keys(data.tree)) {
      if (shownBanks.has(b)) continue;
      const bankName = data.bank_names[b] || b;
      const count = data.bank_counts[b] || 0;
      html += '<div class="wiki-bank" id="wiki-bank-' + b + '">';
      html += '<div class="wiki-bank-header" data-bank="' + b + '">';
      html += '<span class="icon">📁</span>';
      html += '<span class="name">' + bankName + '</span>';
      html += '<span class="count">' + count + ' 文档</span>';
      html += '<span class="arrow">▶</span>';
      html += '</div>';
      html += '<div class="wiki-bank-body">';
      const cats = data.tree[b];
      for (const cat of Object.keys(cats).sort()) {
        const docs = cats[cat];
        html += '<div class="wiki-cat">';
        html += '<div class="wiki-cat-header" data-cat="' + cat + '">';
        html += '<span class="cat-icon">📁</span>';
        html += '<span class="cat-name">' + cat + '</span>';
        html += '<span class="cat-count">' + docs.length + '</span>';
        html += '<span class="cat-arrow">▶</span>';
        html += '</div>';
        html += '<div class="wiki-cat-body">';
        for (const doc of docs) {
          const typeLabel = doc.doc_type === 'gb_standard' ? 'GB标准' : doc.doc_type === 'regulation' ? '法规' : '';
          html += '<div class="wiki-doc" data-docid="' + doc.id + '">';
          html += '<span class="doc-icon">📄</span>';
          html += '<span class="doc-title">' + doc.title.replace(/</g, '&lt;') + '</span>';
          if (typeLabel) html += '<span class="doc-type">' + typeLabel + '</span>';
          html += '</div>';
          html += '<div class="wiki-doc-content" id="wiki-doc-' + doc.id + '"></div>';
        }
        html += '</div></div>';
      }
      html += '</div></div>';
    }
    
    treeDiv.innerHTML = html;
    // Event delegation for wiki tree clicks
    treeDiv.addEventListener('click', function(e) {
      var bankHeader = e.target.closest('.wiki-bank-header[data-bank]');
      if (bankHeader) { toggleWikiBank(bankHeader.getAttribute('data-bank')); return; }
      var catHeader = e.target.closest('.wiki-cat-header[data-cat]');
      if (catHeader) { toggleWikiCat(catHeader.parentElement); return; }
      var docEl = e.target.closest('.wiki-doc[data-docid]');
      if (docEl) { toggleWikiDoc(docEl, docEl.getAttribute('data-docid')); return; }
    });
  } catch(e) {
    treeDiv.innerHTML = '<div class="result-msg error">加载失败: ' + e.message + '</div>';
  }
}

function toggleWikiBank(bankId) {
  const el = document.getElementById('wiki-bank-' + bankId);
  if (el) el.classList.toggle('open');
}

function toggleWikiCat(el) {
  el.classList.toggle('open');
}

async function toggleWikiDoc(el, docId) {
  const contentDiv = document.getElementById('wiki-doc-' + docId);
  if (!contentDiv) return;
  
  if (contentDiv.classList.contains('show')) {
    contentDiv.classList.remove('show');
    return;
  }
  
  contentDiv.classList.add('show');
  contentDiv.innerHTML = '<div class="loading"><span class="spin"></span>加载文档内容...</div>';
  
  try {
    const r = await fetch('/api/documents/' + docId + '/content');
    const d = await r.json();
    if (!d.text) {
      contentDiv.innerHTML = '<div class="result-msg error">' + (d.detail || '内容未找到') + '</div>';
      return;
    }
    const safeText = d.text.replace(/</g, '&lt;');
    contentDiv.innerHTML = '<strong>📄 ' + (d.title || '未知').replace(/</g, '&lt;') + '</strong>（' + d.chunks + ' 个片段，' + d.text.length + ' 字符）<hr style="border-color:#eee;margin:6px 0">' + safeText.substring(0, 5000) + (d.text.length > 5000 ? '<br><em style="color:#888">...（仅显示前5000字符）</em>' : '');
  } catch(e) {
    contentDiv.innerHTML = '<div class="result-msg error">加载失败: ' + e.message + '</div>';
  }
}

async function doWebSearch() {
  const q = document.getElementById('query').value.trim();
  if (!q) return;
  const area = document.getElementById('answer-area');
  
  // 收集当前页面上下文
  const pageContext = area ? area.innerText : '';
  
  // 显示 loading
  const loadingDiv = document.createElement('div');
  loadingDiv.className = 'loading';
  loadingDiv.innerHTML = '<span class="spin"></span> 联网搜索中...';
  area.appendChild(loadingDiv);
  
  try {
    const fd = new FormData();
    fd.append('q', q);
    fd.append('bank', currentBank);
    fd.append('context', pageContext.substring(0, 4000));
    
    const r = await fetch('/api/web-search', {method: 'POST', body: fd});
    const d = await r.json();
    
    loadingDiv.remove();
    
    let html = '<div class="answer" style="border-left:3px solid #4CAF50;margin-top:12px">';
    html += '<span style="display:inline-block;padding:2px 8px;background:#e8f5e9;color:#2e7d32;border-radius:4px;font-size:12px;font-weight:600;margin-bottom:6px">🌐 联网搜索</span><br>';
    html += d.answer.replace(/\\\\\\n/g, '<br>');
    html += '</div>';
    
    area.innerHTML += html;
  } catch(e) {
    loadingDiv.remove();
    area.innerHTML += '<div class="result-msg error">联网搜索失败: ' + e.message + '</div>';
  }
}

async function doFetchStandard() {
  const stdNo = document.getElementById('std-no').value.trim();
  if (!stdNo) return;
  const bank = document.getElementById('std-bank').value || 'kb';
  const resultDiv = document.getElementById('std-result');
  
  resultDiv.innerHTML = '<div class="loading"><span class="spin"></span>搜索中...</div>';
  
  try {
    const fd = new FormData();
    fd.append('std_no', stdNo);
    fd.append('bank', bank);
    
    resultDiv.innerHTML = '<div class="loading"><span class="spin"></span>下载中...</div>';
    const r = await fetch('/api/fetch-standard', {method:'POST', body:fd});
    const d = await r.json();
    
    if (d.ok) {
      const bankName = (bankData.find(b => b.key === d.bank) || {}).name || d.bank;
      resultDiv.innerHTML = '<div class="result-msg success">✅ ' + d.title + ' 已入库<br>' +
        '<small>分类：' + bankName + ' | ' + (d.text_length||0) + ' 字符 | ' + (d.chunks||'?') + ' chunks</small><br>' +
        '<small style="color:#999">doc_id: ' + (d.doc_id||'?').substring(0,8) + '...</small></div>';
      document.getElementById('std-no').value = '';
      loadStats(); loadBanks();
    } else {
      resultDiv.innerHTML = '<div class="result-msg error">' + (d.detail || '下载失败') + '</div>';
    }
  } catch(e) {
    resultDiv.innerHTML = '<div class="result-msg error">' + e.message + '</div>';
  }
}

let uploadQueue = [];  // {file, title}

async function doAudit() {
  const resultDiv = document.getElementById('audit-result');
  const q = String.fromCharCode(39);
  resultDiv.innerHTML = '<div class="loading"><span class="spin"></span>正在扫描知识库...</div>';

  try {
    const r = await fetch('/api/audit');
    const d = await r.json();

    let html = '<div class="result-msg success" style="margin-bottom:12px">';
    html += '📊 共 ' + d.total_docs + ' 篇文档，平均质量 ' + d.avg_score + '%';
    if (d.low_quality_count > 0) {
      html += ' · <span style="color:#c62828">⚠️ ' + d.low_quality_count + ' 篇需关注</span>';
    }
    html += '</div>';

    if (d.documents && d.documents.length) {
      html += '<div style="max-height:500px;overflow-y:auto">';
      d.documents.forEach(doc => {
        const color = doc.score >= 80 ? '#2e7d32' : doc.score >= 60 ? '#e67e22' : '#c62828';
        const badge = doc.needs_refetch ? ' <span style="background:#c62828;color:#fff;padding:1px 6px;border-radius:4px;font-size:10px">需重下</span>' : '';
        const refetchBtn = doc.needs_refetch ? `<button onclick="doRefetch(${q}${doc.doc_id}${q},${q}${(doc.title||'').replace(new RegExp(q,'g'),'\\\\'+q)}${q},this)" style="font-size:10px;padding:2px 6px;background:#e94560;color:#fff;border:none;border-radius:3px;cursor:pointer">🔄 重下</button>` : '';
        html += '<div style="padding:8px 12px;border-bottom:1px solid #f0f0f0;font-size:13px">';
        html += '<span style="font-weight:600;color:' + color + '">' + doc.score + '%</span>';
        html += ' <span>' + doc.title + '</span>' + badge + refetchBtn;
        html += '<div style="color:#999;font-size:11px">' + (doc.chars||0) + ' 字符 · ' + (doc.issues||[]).join(' · ') + '</div>';
        html += '</div>';
      });
      html += '</div>';
    }
    resultDiv.innerHTML = html;
  } catch(e) {
    resultDiv.innerHTML = '<div class="result-msg error">自查失败: ' + e.message + '</div>';
  }
}

async function doRefetch(docId, title, btn) {
  if (!confirm('确认重新下载「' + title + '」？将使用 MinerU 重新解析并替换原有数据。')) return;
  btn.disabled = true;
  btn.textContent = '⏳';
  try {
    const fd = new FormData();
    fd.append('doc_id', docId);
    fd.append('std_no', title);
    const r = await fetch('/api/audit/refetch', {method:'POST', body:fd});
    const d = await r.json();
    if (d.ok) {
      btn.textContent = '✅';
      btn.style.background = '#27ae60';
      alert('重新入库成功！质量评分: ' + d.new_score + '%');
    } else {
      btn.textContent = '❌';
      alert('失败: ' + (d.detail || '未知错误'));
    }
  } catch(e) {
    btn.textContent = '❌';
    alert('失败: ' + e.message);
  }
}

function onFilesSelected(fileList) {
  if (!fileList || fileList.length === 0) return;
  uploadQueue = [];
  const hint = document.getElementById('drop-hint');
  
  for (const f of fileList) {
    uploadQueue.push({file: f, title: f.name.replace(/\.[^.]+$/, '')});
  }
  
  if (uploadQueue.length === 1) {
    const f = uploadQueue[0].file;
    hint.innerHTML = `<span class="file-selected">📎 ${f.name}</span><br><small style="color:#999">${(f.size/1024).toFixed(1)} KB</small>`;
  } else {
    const totalSize = uploadQueue.reduce((s, q) => s + q.file.size, 0);
    hint.innerHTML = `<span class="file-selected">📎 ${uploadQueue.length} 个文件</span><br><small style="color:#999">共 ${(totalSize/1024/1024).toFixed(1)} MB</small>`;
  }
  
  document.getElementById('upload-btn').disabled = false;
  document.getElementById('upload-btn').textContent = `上传 ${uploadQueue.length} 个文件`;
}

async function doUpload() {
  if (uploadQueue.length === 0) return;
  const btn = document.getElementById('upload-btn');
  const resultDiv = document.getElementById('upload-result');
  const category = document.getElementById('upload-category').value;
  const uploadBank = document.getElementById('upload-bank').value || (currentBank === 'all' ? 'kb' : currentBank);
  
  btn.disabled = true;
  
  // Build queue UI
  let html = '<div style="font-size:13px;margin-bottom:8px;color:#555">📋 上传队列</div>';
  uploadQueue.forEach((q, i) => {
    html += `<div id="q-${i}" style="padding:6px 10px;margin:2px 0;background:#f8f8f8;border-radius:4px;font-size:12px;display:flex;align-items:center;gap:8px">
      <span style="color:#bbb">⏳</span>
      <span style="flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap">${q.file.name}</span>
      <span style="color:#999;flex-shrink:0">${(q.file.size/1024).toFixed(0)}KB</span>
    </div>`;
  });
  resultDiv.innerHTML = html;
  
  // Process queue
  let ok = 0, fail = 0, skip = 0;
  for (let i = 0; i < uploadQueue.length; i++) {
    const q = uploadQueue[i];
    const row = document.getElementById('q-' + i);
    row.querySelector('span').textContent = '⏺';
    row.style.background = '#fff8e1';
    
    const fd = new FormData();
    fd.append('file', q.file);
    fd.append('title', q.title);
    fd.append('category', category);
    fd.append('bank', uploadBank);
    
    try {
      const r = await fetch('/api/upload', {method:'POST', body:fd});
      const d = await r.json();
      if (r.status === 409) {
        row.querySelector('span').textContent = '⚠️';
        row.style.background = '#fff3e0';
        row.title = d.detail || '文档已存在';
        skip++;
      } else if (d.quality && d.quality.needs_confirm && !d.ok) {
        row.querySelector('span').textContent = '⚠️';
        row.style.background = '#fff3e0';
        row.title = '质量评分: ' + d.quality.score + '% - ' + (d.quality.issues || []).join(', ');
        const NL = String.fromCharCode(10);
        const confirmed = confirm(
          '文档解析质量较低（' + d.quality.score + '%）' + NL + NL +
          '问题: ' + (d.quality.issues || []).join(NL) + NL + NL +
          '是否仍然上传？建议取消后使用 MinerU 重新解析。'
        );
        if (confirmed) {
          const fd2 = new FormData();
          fd2.append('file', q.file);
          fd2.append('title', q.title);
          fd2.append('category', category);
          fd2.append('bank', uploadBank);
          fd2.append('confirm_quality', 'true');
          const r2 = await fetch('/api/upload', {method:'POST', body:fd2});
          const d2 = await r2.json();
          if (d2.ok) {
            row.querySelector('span').textContent = '✅';
            row.style.background = '#e8f5e9';
            ok++;
          } else {
            row.querySelector('span').textContent = '❌';
            row.style.background = '#ffebee';
            fail++;
          }
        } else {
          skip++;
        }
      } else if (d.ok) {
        row.querySelector('span').textContent = '✅';
        row.style.background = '#e8f5e9';
        let tip = '';
        if (d.quality) tip += '质量评分: ' + d.quality.score + '%';
        if (d.integrity && d.integrity.status === 'ok') tip += ' · 完整性: ' + d.integrity.coverage_pct + '%';
        row.title = tip;
        ok++;
      } else {
        row.querySelector('span').textContent = '❌';
        row.style.background = '#ffebee';
        row.title = d.detail || '上传失败';
        fail++;
      }
    } catch(e) {
      row.querySelector('span').textContent = '❌';
      row.style.background = '#ffebee';
      row.title = e.message;
      fail++;
    }
  }
  
  // Summary
  const summary = [];
  if (ok > 0) summary.push(`✅ ${ok} 成功`);
  if (skip > 0) summary.push(`⚠️ ${skip} 跳过`);
  if (fail > 0) summary.push(`❌ ${fail} 失败`);
  resultDiv.insertAdjacentHTML('afterbegin', 
    `<div class="result-msg success" style="margin-bottom:8px">${summary.join(' · ')}</div>`);
  
  uploadQueue = [];
  btn.textContent = '上传到知识库';
  document.getElementById('drop-hint').innerHTML = '点击选择文件 / 拖拽上传<br>支持 PDF / Word / Markdown / TXT<br><small style="color:#aaa">支持多选和文件夹</small>';
  loadStats(); loadBanks();
}

async function loadDocs() {
  const list = document.getElementById('doc-list');
  list.innerHTML = '<div class="loading"><span class="spin"></span>加载中...</div>';
  try {
    const r = await fetch('/api/documents?bank=' + currentBank);
    const d = await r.json();
    allDocs = d.documents;
    renderDocs();
  } catch(e) {
    list.innerHTML = '<div class="result-msg error">加载失败</div>';
  }
}

function renderDocs() {
  const list = document.getElementById('doc-list');
  if (!allDocs.length) {
    list.innerHTML = '<div class="empty">暂无文档</div>'; return;
  }
  const q = (document.getElementById('doc-search').value || '').toLowerCase();
  const filtered = q ? allDocs.filter(d => d.title.toLowerCase().includes(q) || (d.filename||'').toLowerCase().includes(q)) : allDocs;
  if (!filtered.length) {
    list.innerHTML = '<div class="empty">无匹配文档</div>'; return;
  }
  list.innerHTML = filtered.map(doc => {
    const catHtml = doc.category ? `<span class="cat-tag">${doc.category}</span>` : '';
    const dateStr = doc.created ? new Date(doc.created).toLocaleDateString('zh-CN') : '-';
    const bankOpts = bankData.filter(b => b.key !== 'all').map(b =>
      `<option value="${b.key}"${doc.bank === b.key ? ' selected' : ''}>${b.name}</option>`
    ).join('');
    return `<div class="doc-item" id="doc-row-${doc.id}">
      <div class="main">
        <div class="title-line">
          <span class="title">📄 ${doc.title}</span>${catHtml}
        </div>
        <div class="meta">${doc.chunks > 0 ? doc.chunks + ' 片段 · ' + doc.size_chars + ' 字符' : '—'} · ${dateStr}</div>
      </div>
      <div class="actions">
        <select class="doc-bank-select" onchange="changeDocBank('${doc.id}', this.value)" title="切换分类">${bankOpts}</select>
        <span class="del" onclick="delDoc('${doc.id}',this)">🗑</span>
      </div>
    </div>`;
  }).join('');
}

function populateDocBankSelects() {
  document.querySelectorAll('.doc-bank-select').forEach(sel => {
    const currentVal = sel.value;
    sel.innerHTML = bankData.filter(b => b.key !== 'all').map(b =>
      `<option value="${b.key}"${b.key === currentVal ? ' selected' : ''}>${b.name}</option>`
    ).join('');
  });
}

async function changeDocBank(docId, newBank) {
  if (!newBank || newBank === 'all') return;
  try {
    const fd = new FormData(); fd.append('bank', newBank);
    const r = await fetch('/api/documents/' + docId + '/bank', {method:'PATCH', body:fd});
    if (r.ok) {
      // Update local data
      const doc = allDocs.find(d => d.id === docId);
      if (doc) doc.bank = newBank;
      loadBanks();
    }
  } catch(e) {}
}

async function delDoc(id, el) {
  if (!confirm('确认删除此文档及所有关联内容？')) return;
  await fetch('/api/documents/'+id, {method:'DELETE'});
  el.closest('.doc-item').remove();
  allDocs = allDocs.filter(d => d.id !== id);
  loadStats(); loadBanks();
}

let catData = [];
async function loadCategories() {
  try {
    const r = await fetch('/api/categories');
    catData = (await r.json()).categories;
    const sel = document.getElementById('upload-category');
    if (sel.options.length <= 1) {
      catData.forEach(c => { if (c.name !== '其他') sel.appendChild(new Option(c.name, c.name)); });
      sel.appendChild(new Option('其他', '其他'));
    }
  } catch(e) {}
}

async function loadStats() {
  try {
    const r = await fetch('/api/stats');
    const d = await r.json();
    document.getElementById('stats').innerHTML = `<span>📄 ${d.total_documents||0} 文档</span><span>🧩 ${d.total_nodes||0} 节点</span>`;
  } catch(e) {}
}

const dropZone = document.getElementById('drop-zone');
dropZone.addEventListener('dragover', e => { e.preventDefault(); dropZone.classList.add('dragover'); });
dropZone.addEventListener('dragleave', () => dropZone.classList.remove('dragover'));
dropZone.addEventListener('drop', e => {
  e.preventDefault(); dropZone.classList.remove('dragover');
  if (e.dataTransfer.files.length > 0) onFilesSelected(e.dataTransfer.files);
});
</script>
</body>
</html>"""


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=3002)
